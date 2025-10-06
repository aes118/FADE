# Grant Application Portal – Logframe + Workplan + Budget
import streamlit as st
import pandas as pd
from io import BytesIO
from openpyxl import Workbook
import uuid
import base64
import os
import html, re
from contextlib import contextmanager
from html import escape
from datetime import datetime, date
import hashlib
import streamlit as st, requests, tempfile, os
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.shared import Pt
from docx.oxml.ns import qn

# ---------------- Page config ----------------
st.set_page_config(page_title="Falcon Awards Project Portal", layout="wide")
st.sidebar.image("glide_logo.png", width="stretch")

# ---------------- Canonical field labels ----------------
LABELS = {
    "title": "Project title",
    "pi_name": "Principal Investigator (PI) name",
    "pi_email": "PI email",
    "implementing_partners": "Implementing Partner(s)",
    "supporting_partners": "Supporting Partners (Optional)",
    "start_date": "Project start date",
    "end_date": "Project end date",
    "location": "Implementation location",
    "contact_name": "Main Contact person (Optional)",
    "contact_email": "Main Contact email (Optional)",
    "contact_phone": "Contact phone (Optional)",
    "total_funding": "Total funding requested (From Budget)",
    "outputs_count": "# Outputs (From Logframe)",
    "kpis_count": "# KPIs (From Logframe)",
    "activities_count": "# Activities (From Workplan)",
}

# ---------------- Helpers & State ----------------
def _s(v):
    try:
        import pandas as pd
        if pd.isna(v):
            return ""
    except Exception:
        pass
    return str(v).strip()

def generate_id():
    return str(uuid.uuid4())[:8]

footer_note = "Note: All monetary values are in USD and will be converted to AED in the official project agreement."

# app state
for key in ["impacts", "outcomes", "outputs", "kpis", "workplan", "budget", "disbursement"]:
    if key not in st.session_state:
        st.session_state[key] = []

# edit-state for inline editing
for key in ["edit_goal", "edit_outcome", "edit_output", "edit_kpi", "edit_activity", "edit_budget_row"]:
    if key not in st.session_state:
        st.session_state[key] = None

def _find_by_id(lst, _id):
    for i, x in enumerate(lst):
        if x.get("id") == _id:
            return i
    return None

def delete_cascade(*, goal_id=None, outcome_id=None, output_id=None):
    """Delete an item and its children in the Goal->Outcome->Output->KPI hierarchy."""
    if goal_id:
        for oc in [o for o in st.session_state.outcomes if o.get("parent_id") == goal_id]:
            delete_cascade(outcome_id=oc["id"])
        st.session_state.impacts = [g for g in st.session_state.impacts if g["id"] != goal_id]

    if outcome_id:
        for out in [o for o in st.session_state.outputs if o.get("parent_id") == outcome_id]:
            delete_cascade(output_id=out["id"])
        st.session_state.outcomes = [o for o in st.session_state.outcomes if o["id"] != outcome_id]

    if output_id:
        # delete KPIs under output
        st.session_state.kpis = [k for k in st.session_state.kpis if k.get("parent_id") != output_id]
        st.session_state.outputs = [o for o in st.session_state.outputs if o["id"] != output_id]

def render_editable_item(
    *,
    item: dict,
    list_name: str,
    edit_flag_key: str,
    view_md_func,
    fields=None,
    default_label="Name",
    on_delete=None,
    key_prefix="logframe"   # <— NEW
):
    c1, c2, c3 = st.columns([0.85, 0.07, 0.08])

    # unique widget id base
    wid = f"{key_prefix}_{list_name}_{item['id']}"

    if st.session_state.get(edit_flag_key) == item["id"]:
        new_values = {}
        if fields:
            for fkey, widget_func, label in fields:
                new_values[fkey] = widget_func(label, value=item.get(fkey, ""), key=f"{wid}_{fkey}")
        else:
            new_values["name"] = c1.text_input(default_label, value=item.get("name", ""), key=f"{wid}_name")

        if c2.button("💾", key=f"{wid}_save"):
            idx = _find_by_id(st.session_state[list_name], item["id"])
            if idx is not None:
                # Defensive: prevent users from saving names that include labels
                if "name" in new_values:
                    if list_name == "activities":
                        new_values["name"] = strip_label_prefix(new_values["name"], "Activity")
                    elif list_name == "kpis":
                        new_values["name"] = strip_label_prefix(new_values["name"], "KPI")

                for k, v in new_values.items():
                    st.session_state[list_name][idx][k] = v.strip() if isinstance(v, str) else v
            st.session_state[edit_flag_key] = None
            st.rerun()

        if c3.button("✖️", key=f"{wid}_cancel"):
            st.session_state[edit_flag_key] = None
            st.rerun()
    else:
        c1.markdown(view_md_func(item), unsafe_allow_html=True)
        if c2.button("✏️", key=f"{wid}_edit"):
            st.session_state[edit_flag_key] = item["id"]
            st.rerun()
        if c3.button("🗑️", key=f"{wid}_del"):
            if on_delete:
                on_delete()

from datetime import date

def compute_numbers(include_activities: bool = False):
    """
    Preserve user/excel order.
    Outputs: numbered per Outcome in list order.
    KPIs:    numbered per Output in list order (as stored in st.session_state.kpis).
    Activities (if requested): numbered per Output in list order.
    """
    out_num, kpi_num = {}, {}
    outcomes = st.session_state.get("outcomes", [])
    outputs  = st.session_state.get("outputs", [])
    kpis     = st.session_state.get("kpis", [])

    # Outputs numbered per Outcome (list order)
    for oc in outcomes:
        oc_outs = [o for o in outputs if o.get("parent_id") == oc["id"]]  # list order preserved
        for i, out in enumerate(oc_outs, start=1):
            out_num[out["id"]] = f"{i}"

    # KPIs numbered per Output (list order)
    for out_id, n in out_num.items():
        p = 1
        for k in kpis:                       # iterate as-is → preserves entry/import order
            if k.get("parent_id") == out_id:
                kpi_num[k["id"]] = f"{n}.{p}"
                p += 1

    if not include_activities:
        return out_num, kpi_num

    # Activities numbered per Output (list order)
    act_num = {}
    workplan = st.session_state.get("workplan", [])
    for out_id, n in out_num.items():
        q = 1
        for a in workplan:
            if a.get("output_id") == out_id:
                act_num[a["id"]] = f"{n}.{q}"
                q += 1
    return out_num, kpi_num, act_num


def format_activity_label(activity: dict, act_nums: dict, *, unknown_placeholder: str = "?") -> str:
    """Return the standard display label for an activity."""
    if not activity:
        return ""

    act_id = activity.get("id")
    name = activity.get("name", "")
    number = act_nums.get(act_id)

    if not number:
        number = unknown_placeholder or ""

    prefix = f"Activity {number}".strip()
    if not prefix:
        prefix = "Activity"

    if name:
        return f"{prefix} — {name}"
    return prefix


def activity_label_map(act_nums=None, *, unknown_placeholder: str = "?") -> dict:
    """Build a {activity_id -> label} mapping reused across tables/exports."""
    if act_nums is None:
        _, _, act_nums = compute_numbers(include_activities=True)

    return {
        a["id"]: format_activity_label(a, act_nums, unknown_placeholder=unknown_placeholder)
        for a in st.session_state.get("workplan", [])
    }


def _workplan_df():
    """Return a tidy DF with Activity/Output labels and metadata (rows that have both dates)."""
    import pandas as pd

    outs = {o["id"]: (o.get("name") or "Output") for o in st.session_state.get("outputs", [])}
    out_nums, _, act_nums = compute_numbers(include_activities=True)

    rows = []
    for a in st.session_state.get("workplan", []):
        s, e = a.get("start"), a.get("end")
        if not (s and e):
            continue

        out_id = a.get("output_id")
        out_name = outs.get(out_id, "(unassigned)")
        out_num = out_nums.get(out_id, "")
        if out_num and out_name:
            out_label = f"Output {out_num} — {out_name}"
        elif out_num:
            out_label = f"Output {out_num}"
        else:
            out_label = out_name or "Output"

        act_name = a.get("name") or ""
        act_num = act_nums.get(a.get("id"), "")
        if act_num and act_name:
            act_label = f"Activity {act_num} — {act_name}"
        elif act_num:
            act_label = f"Activity {act_num}"
        else:
            act_label = act_name or "Activity"

        owner = a.get("owner")
        if owner is None:
            owner = ""

        rows.append({
            "Activity": act_name,
            "ActivityLabel": act_label,
            "ActivityID": a.get("id"),
            "Output": out_name,
            "OutputLabel": out_label,
            "OutputID": out_id,
            "Owner": owner,
            "Start": s,
            "End": e,
        })

    if not rows:
        return pd.DataFrame(
            columns=[
                "Activity",
                "ActivityLabel",
                "ActivityID",
                "Output",
                "OutputLabel",
                "OutputID",
                "Owner",
                "Start",
                "End",
            ]
        )

    return (
        pd.DataFrame(rows)
        .sort_values(["Output", "Start", "ActivityLabel"], kind="stable")
        .reset_index(drop=True)
    )

def _draw_gantt(ax, df, *, show_today=False):
    """
    Clean, readable Gantt:
    - Bars colored by Output with subtle edges
    - Group separators
    - Major ticks by quarter, minor ticks monthly; labels MMM-YYYY
    - Optional 'today' line
    - Bottom-centered legend
    """
    import numpy as np
    import matplotlib.pyplot as plt
    import matplotlib.dates as mdates
    from matplotlib.patches import Patch
    from datetime import date, timedelta

    if df.empty:
        ax.text(0.5, 0.5, "No activities with both start & end dates",
                ha="center", va="center")
        ax.axis("off")
        return

    if "ActivityLabel" not in df.columns:
        df = df.assign(ActivityLabel=df["Activity"])
    if "OutputLabel" not in df.columns:
        df = df.assign(OutputLabel=df["Output"])

    # ---- Order outputs by earliest start for a logical reading order
    first_start = df.groupby("Output")["Start"].min().sort_values()
    ordered_outputs = first_start.index.tolist()
    output_labels = {}
    if "OutputLabel" in df.columns:
        for out, lbl in zip(df["Output"], df["OutputLabel"]):
            output_labels.setdefault(out, lbl)
    df = (
        df.assign(_out_order=df["Output"].map({o:i for i, o in enumerate(ordered_outputs)}))
          .sort_values(["_out_order", "Start", "ActivityLabel"], kind="stable")
          .drop(columns="_out_order")
          .reset_index(drop=True)
    )

    # ---- Stable color mapping (tab20_cycle)
    cmap = plt.get_cmap("tab20")
    color_map = {o: cmap(i % 20) for i, o in enumerate(ordered_outputs)}

    # ---- Build bar geometry
    y      = np.arange(len(df))
    widths = [(e - s).days or 0.5 for s, e in zip(df["Start"], df["End"])]
    lefts  = df["Start"].tolist()
    colors = [color_map[o] for o in df["Output"]]

    # ---- Bars (with gentle edge for contrast)
    ax.barh(
        y, widths, left=lefts, color=colors,
        edgecolor="white", linewidth=0.8, alpha=0.95
    )

    # Optional: label durations at the end of each bar (uncomment to enable)
    # for yi, s, w in zip(y, lefts, widths):
    #     ax.text(s + timedelta(days=w) , yi, f" {int(w)}d",
    #             va="center", ha="left", fontsize=8, color="#444")

    # ---- Y axis: activities
    ax.set_yticks(y)
    ax.set_yticklabels(df["ActivityLabel"], fontsize=9)
    ax.invert_yaxis()
    ax.set_ylabel("Activity")

    # ---- X axis: quarterly major ticks + monthly minors; labels MMM-YYYY
    ax.xaxis.set_major_locator(mdates.MonthLocator(bymonth=(1,4,7,10)))  # Jan/Apr/Jul/Oct
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%b-%Y"))
    ax.xaxis.set_minor_locator(mdates.MonthLocator())  # every month
    for lab in ax.get_xticklabels():
        lab.set_rotation(30); lab.set_ha("right")

    # ---- Range padding (7% of span on both sides)
    x_min = min(df["Start"]); x_max = max(df["End"])
    span  = (x_max - x_min).days or 1
    pad   = max(int(span * 0.07), 7)
    ax.set_xlim(x_min - timedelta(days=pad), x_max + timedelta(days=pad))

    # ---- Grid (minor vertical dots)
    ax.grid(axis="x", which="minor", linestyle=":", linewidth=0.6, alpha=0.35)
    ax.grid(axis="x", which="major", linestyle="-", linewidth=0.6, alpha=0.35)
    ax.set_xlabel("Date")

    # ---- Group separators (between Outputs)
    sep_rows = [i for i in range(1, len(df)) if df.loc[i, "Output"] != df.loc[i-1, "Output"]]
    xmin, xmax = ax.get_xlim()
    for i in sep_rows:
        ax.hlines(i - 0.5, xmin, xmax, colors="#BFC5CC", linestyles="-", linewidth=0.8, alpha=0.7)
    ax.set_xlim(xmin, xmax)  # keep full-width lines

    # ---- Optional 'today' line
    if show_today:
        today = date.today()
        ax.axvline(today, color="#666", linewidth=1.1, alpha=0.8)
        ax.text(today, ax.get_ylim()[0] - 0.25, "Today", rotation=90,
                va="top", ha="center", fontsize=8, color="#666")

    # ---- Legend (bottom center), tidy spacing
    handles = [
        Patch(facecolor=color_map[o], label=output_labels.get(o, o))
        for o in ordered_outputs
    ]
    ax.legend(
        handles=handles,
        title="Outputs",
        loc="upper center",
        bbox_to_anchor=(0.2, -0.30),
        ncol=2,  # wrap if many outputs
        # ncol=min(len(handles), 4),  # wrap if many outputs
        frameon=False,
        handlelength=1.8,
        columnspacing=1.4,
        handletextpad=0.6,
        borderaxespad=0.0,
    )

def select_output_id(label, current_id, key):
    """Select an Output and return its ID (options are IDs; stable + clean)."""
    outs = st.session_state.get("outputs", [])
    out_nums, _ = compute_numbers()
    id_to_name = {o["id"]: (o.get("name") or "Output") for o in outs}
    options = [o["id"] for o in outs]
    idx = options.index(current_id) if current_id in options else 0
    return st.selectbox(
        label,
        options,
        index=idx if options else 0,
        format_func=lambda oid: f"Output {out_nums.get(oid,'?')} — {id_to_name.get(oid,'Output')}",
        key=key,
    )

def strip_label_prefix(text: str, kind: str) -> str:
    """
    Remove labels like 'Activity 1.2 — ' or 'KPI 1.2.3: ' from a string.
    Accepts separators '—', ':', or '-'.
    """
    if not isinstance(text, str):
        return text
    pat = rf'^\s*{kind}\s+\d+(?:\.\d+)*\s*[—:\-]\s*'
    return re.sub(pat, '', text).strip()

def parse_date_like(v):
    """Return a datetime.date or None from common date formats or existing date/datetime/pandas types."""
    if v is None:
        return None

    # Handle pandas NaT / NaN early
    try:
        import pandas as pd
        if pd.isna(v):
            return None
        if isinstance(v, pd.Timestamp):
            return v.date()
    except Exception:
        pass

    if isinstance(v, datetime):
        return v.date()
    if isinstance(v, date):
        return v

    # Strings
    s = str(v).strip()
    if not s or s.lower() in ("none", "nan", "nat"):
        return None

    # Try a bunch of common formats (with and without HH:MM:SS)
    fmts = (
        "%Y-%m-%d",
        "%Y-%m-%d %H:%M:%S",
        "%d/%m/%Y",
        "%d/%m/%Y %H:%M:%S",
        "%m/%d/%Y",
        "%m/%d/%Y %H:%M:%S",
        "%Y/%m/%d",
        "%Y/%m/%d %H:%M:%S",
        "%d/%b/%Y",            # 03/Sep/2025
        "%d/%b/%Y %H:%M:%S",
        "%d-%b-%Y",            # 03-Sep-2025
        "%d-%b-%Y %H:%M:%S",
    )
    for fmt in fmts:
        try:
            return datetime.strptime(s, fmt).date()
        except ValueError:
            continue

    return None

def fmt_dd_mmm_yyyy(v):
    """Return 'DD/MMM/YYYY' (e.g., 03/Sep/2025) or '' if not set/parsable."""
    try:
        import pandas as pd
        if pd.isna(v):                   # handles NaT / NaN
            return ""
        if isinstance(v, pd.Timestamp):  # valid timestamp -> format directly
            return v.strftime("%d/%b/%Y")
    except Exception:
        pass

    d = parse_date_like(v)
    return d.strftime("%d/%b/%Y") if d else ""

def fmt_money(val) -> str:
    """Return number with thousands dot and 2 decimals, e.g., 1.234.567,89."""
    try:
        x = float(val)
    except (TypeError, ValueError):
        return ""
    # First format in US style, then swap separators
    s = f"{x:,.2f}"                  # -> 1,234,567.89
    return s.replace(",", "␟").replace(".", ",").replace("␟", ".")

def view_logframe_element(inner_html: str, kind: str = "output") -> str:
    """Wrap inner HTML in a styled card. kind: 'output' | 'kpi' (or others later)."""
    return f"<div class='lf-card lf-card--{kind}'>{inner_html}</div>"


def _assumption_items(text: str) -> list:
    """Return cleaned bullet items from a multiline assumptions string."""
    items = []
    for ln in (text or "").splitlines():
        cleaned = re.sub(r"^[\-\u2022]\s*", "", str(ln)).strip()
        if cleaned:
            items.append(cleaned)
    return items


def _assumptions_html(text: str) -> str:
    """HTML block (heading + list) for assumptions."""
    items = _assumption_items(text)
    if not items:
        return ""
    lis = "".join(f"<li>{html.escape(x)}</li>" for x in items)
    return (
        "<div class='lf-ass'>"
        "<div class='lf-ass-heading'><b> Key Assumptions </b></div>"
        f"<ul class='lf-ass-list'>{lis}</ul>"
        "</div>"
    )


def _assumptions_doc_text(text: str) -> str:
    """Plain text bullet list (with •) for Word export."""
    items = _assumption_items(text)
    return "\n".join(f"• {item}" for item in items)

def sync_disbursement_from_kpis():
    """
    Keep st.session_state.disbursement in sync with KPIs:
    - rows only for KPIs with linked_payment == True
    - preserve user-entered anticipated_date and amount_usd
    - refresh output_id and kpi_name from KPI
    - if date is missing, prefill with latest linked Activity end, else KPI end, else KPI start
    """
    pay_kpis = [k for k in st.session_state.kpis if bool(k.get("linked_payment"))]
    by_id = {d.get("kpi_id"): d for d in st.session_state.disbursement}
    keep_ids = set()

    # helper: latest activity end for a KPI
    def latest_activity_end(kpi_id):
        dates = []
        for a in st.session_state.get("workplan", []):
            if kpi_id in (a.get("kpi_ids") or []):
                if a.get("end"):
                    dates.append(a["end"])
        return max(dates) if dates else None

    for k in pay_kpis:
        kid = k["id"]
        row = by_id.get(kid)
        lae = latest_activity_end(kid)
        if row is None:
            st.session_state.disbursement.append({
                "kpi_id": kid,
                "output_id": k.get("parent_id"),
                "kpi_name": k.get("name", ""),
                "anticipated_date": lae or k.get("end_date") or k.get("start_date") or None,
                "deliverable": k.get("name", ""),
                "amount_usd": 0.0,
            })
        else:
            # refresh mirror fields
            row["output_id"] = k.get("parent_id")
            row["kpi_name"]  = k.get("name", "")
            # if user hasn't chosen a date yet, backfill one
            if not row.get("anticipated_date"):
                row["anticipated_date"] = lae or k.get("end_date") or k.get("start_date") or None
        keep_ids.add(kid)

    # remove rows for KPIs no longer payment-linked
    st.session_state.disbursement = [d for d in st.session_state.disbursement if d.get("kpi_id") in keep_ids]

def _force_calibri_everywhere(doc):
    # 1. Update the document's default styles
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3", "Table Grid"]:
        try:
            style = doc.styles[style_name]
            font = style.font
            font.name = "Calibri"
            font.size = Pt(10)
            rpr = style.element.rPr
            if rpr is not None:
                rFonts = rpr.rFonts
                if rFonts is not None:
                    rFonts.set(qn("w:ascii"), "Calibri")
                    rFonts.set(qn("w:hAnsi"), "Calibri")
                    rFonts.set(qn("w:cs"), "Calibri")
        except KeyError:
            # style may not exist in this template
            pass

    # 2. Sweep through all paragraphs and runs
    for p in doc.paragraphs:
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            rpr = run._element.rPr
            if rpr is not None:
                rFonts = rpr.rFonts
                if rFonts is not None:
                    rFonts.set(qn("w:ascii"), "Calibri")
                    rFonts.set(qn("w:hAnsi"), "Calibri")
                    rFonts.set(qn("w:cs"), "Calibri")

    # 3. Sweep through all table cells
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = "Calibri"
                        run.font.size = Pt(10)
                        rpr = run._element.rPr
                        if rpr is not None:
                            rFonts = rpr.rFonts
                            if rFonts is not None:
                                rFonts.set(qn("w:ascii"), "Calibri")
                                rFonts.set(qn("w:hAnsi"), "Calibri")
                                rFonts.set(qn("w:cs"), "Calibri")

def _render_budget_row_editor(rec: dict, rid: str, act_lookup: dict) -> None:
    """Standalone editor card for a single budget row (rid).
    - Seeds once from rec; preserves user edits across reruns
    - No value=/index=; session_state is the single source of truth
    - Icons only (💾 / ✖️); total chip at the bottom-right
    """
    ss = st.session_state

    # ---------- keys ----------
    init_key      = f"e_init_{rid}"
    item_key      = f"e_item_{rid}"
    act_key       = f"e_act_id_{rid}"
    cat_key       = f"e_cat_{rid}"
    sub_key       = f"e_sub_{rid}"
    unit_choice   = f"e_unit_choice_{rid}"
    unit_custom   = f"e_unit_custom_{rid}"
    uc_key        = f"e_uc_{rid}"
    qty_key       = f"e_qty_{rid}"

    # ---------- 1) Seed once when editor opens ----------
    if not ss.get(init_key, False):
        ss[item_key] = rec.get("item", "")
        ss[act_key]  = rec.get("activity_id")

        ss[cat_key]  = rec.get("category") or list(CATEGORY_TREE.keys())[0]
        sub_opts     = subcategories_for(ss[cat_key]) or ["(none)"]
        existing_sub = rec.get("subcategory")
        ss[sub_key]  = existing_sub if existing_sub in sub_opts else (sub_opts[0] if sub_opts else "")

        # Unit seeding: keep exact match; otherwise 'Custom…' only if needed and allowed
        units, required = unit_choices_for(ss[cat_key], ss[sub_key])
        u_opts = ["Select unit"] + list(dict.fromkeys(units))
        if not required:
            u_opts.append("Custom…")
        existing_unit = (rec.get("unit") or "").strip()
        if existing_unit in u_opts:
            ss[unit_choice] = existing_unit
            ss[unit_custom] = ""
        else:
            if existing_unit and not required:
                ss[unit_choice] = "Custom…"
                ss[unit_custom] = existing_unit
            else:
                ss[unit_choice] = "Select unit"
                ss[unit_custom] = ""

        ss[uc_key]  = float(rec.get("unit_cost") or 0.0)
        ss[qty_key] = float(rec.get("qty") or 0.0)
        ss[init_key] = True

    # ---------- 2) Editor UI (no value=/index=) ----------
    st.markdown("<div class='form-card'>", unsafe_allow_html=True)

    # top-right icons only (save/cancel)
    _, _, ico_save, ico_cancel = st.columns([0.76, 0.14, 0.05, 0.05])
    save_clicked   = ico_save.button("💾", key=f"e_save_icon_{rid}", help="Save")
    cancel_clicked = ico_cancel.button("✖️", key=f"e_cancel_icon_{rid}", help="Cancel")

    # Line Item
    st.text_area("Line Item*", key=item_key, height=48)

    # Linked Activity (optional) — ID-based
    act_options = [None] + list(act_lookup.keys())
    def _fmt_act(aid): return "(none)" if aid is None else act_lookup.get(aid, "(unavailable)")
    st.selectbox("Linked Activity (optional)", act_options, key=act_key, format_func=_fmt_act)

    # Category / Sub-category (stabilized chain)
    c_col, s_col = st.columns(2)
    cur_cat = c_col.selectbox("Cost Category*", list(CATEGORY_TREE.keys()), key=cat_key)

    sub_opts = subcategories_for(cur_cat) or ["(none)"]
    if ss[sub_key] not in sub_opts:
        ss[sub_key] = sub_opts[0]
    cur_sub = s_col.selectbox("Sub Category*", sub_opts, key=sub_key)

    # Unit + custom — recompute after cat/sub changes
    units, required = unit_choices_for(cur_cat, cur_sub)
    u_opts = ["Select unit"] + list(dict.fromkeys(units))
    if not required:
        u_opts.append("Custom…")

    # If current choice became invalid due to cat/sub change, reset gracefully
    if ss[unit_choice] not in u_opts and not (ss[unit_choice] == "Custom…" and not required):
        # keep previously typed custom only if custom is still allowed
        if not required and ss.get(unit_custom, ""):
            ss[unit_choice] = "Custom…"
        else:
            ss[unit_choice] = "Select unit"
            ss[unit_custom] = ""

    st.selectbox("Unit" + ("*" if required else ""), u_opts, key=unit_choice)
    if ss[unit_choice] == "Custom…" and not required:
        st.text_input(" ", key=unit_custom)  # custom unit text
    else:
        ss[unit_custom] = ""

    # Costs
    c1, c2, c3 = st.columns([0.33, 0.33, 0.34])
    st.number_input("Unit Cost (USD)*", min_value=0.0, step=10.0, format="%.2f", key=uc_key)
    st.number_input("Quantity*",       min_value=0.0, step=1.0,  format="%.2f", key=qty_key)

    # bottom-right total chip (end of card)
    new_total = float(ss[uc_key]) * float(ss[qty_key])
    st.markdown(
        f"""
        <div style="display:flex; justify-content:flex-end; margin-top:10px;">
          <span class="total-chip">USD {new_total:,.2f}</span>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.markdown("</div>", unsafe_allow_html=True)  # close form-card

    # ---------- 3) Handlers ----------
    if cancel_clicked:
        ss.edit_budget_row = None
        st.rerun()

    if save_clicked:
        # finalize unit
        choice = ss[unit_choice]
        if choice == "Select unit":
            final_unit = ""
        elif choice == "Custom…" and not required:
            final_unit = (ss.get(unit_custom, "") or "").strip()
        else:
            final_unit = choice

        # validations
        if not ss[item_key].strip():
            st.warning("Line Item is required."); return
        if not cur_cat:
            st.warning("Cost Category is required."); return
        if not cur_sub or cur_sub == "(none)":
            st.warning("Sub Category is required."); return
        if required and not final_unit:
            st.warning("Please choose a Unit for this Sub Category."); return
        if float(ss[uc_key]) <= 0 or float(ss[qty_key]) <= 0:
            st.warning("Unit cost and quantity must be positive."); return

        # update row in place
        rec.update({
            "activity_id": ss[act_key],
            "item": ss[item_key].strip(),
            "category": cur_cat,
            "subcategory": cur_sub,
            "unit": final_unit,
            "unit_cost": float(ss[uc_key]),
            "qty": float(ss[qty_key]),
            "total_usd": new_total,
        })

        ss.edit_budget_row = None
        st.rerun()

def render_pdd(context=None, gantt_image_path: str | None = None):
    """Build the Project Design Document using the branding template with dynamic sections."""
    try:
        from docx import Document
        from docx.shared import Cm, RGBColor, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.section import WD_ORIENT, WD_SECTION_START
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except Exception:
        st.error("`python-docx` is required. In your venv run:\n  pip uninstall -y docx\n  pip install -U python-docx")
        raise

    PRIMARY_SHADE = "0A2F41"

    from io import BytesIO
    from pathlib import Path

    template_path = Path("templates/pdd_template.docx")
    if not template_path.exists():
        st.error(f"Template not found: {template_path}")
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(str(template_path))

    from copy import deepcopy

    body = doc.element.body

    # Keep the section properties (<w:sectPr>) before clearing
    sectPr_el = None
    for child in list(body):
        if child.tag.endswith('sectPr'):
            sectPr_el = deepcopy(child)
        body.remove(child)

    # Reattach the section (or create a new one) so doc.sections[0] is valid
    if sectPr_el is not None:
        body.append(sectPr_el)

    # Safety net: if no sections, add one now
    if len(doc.sections) == 0:
        doc.add_section(WD_SECTION_START.NEW_PAGE)

    def _set_orientation(section, orientation):
        section.orientation = orientation
        if orientation == WD_ORIENT.LANDSCAPE and section.page_width < section.page_height:
            section.page_width, section.page_height = section.page_height, section.page_width
        if orientation == WD_ORIENT.PORTRAIT and section.page_width > section.page_height:
            section.page_width, section.page_height = section.page_height, section.page_width

    # Now it's safe to access and set orientation
    first_section = doc.sections[0]
    _set_orientation(first_section, WD_ORIENT.PORTRAIT)

    def _shade(cell, hex_fill):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_fill)
        tcPr.append(shd)

    def _repeat_header(row):
        trPr = row._tr.get_or_add_trPr()
        tblHeader = OxmlElement("w:tblHeader")
        trPr.append(tblHeader)

    def _set_cell_text(cell, text, *, bold=False, white=False, align_left=True):
        cell.text = ""
        paragraph = cell.paragraphs[0]
        if align_left:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(text or "")
        run.bold = bool(bold)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        try:
            rpr = run._element.rPr
            rpr.rFonts.set(qn("w:ascii"), "Calibri")
            rpr.rFonts.set(qn("w:hAnsi"), "Calibri")
            rpr.rFonts.set(qn("w:cs"), "Calibri")
        except Exception:
            pass
        if white:
            run.font.color.rgb = RGBColor(255, 255, 255)

    def _add_run(paragraph, text, bold=False):
        run = paragraph.add_run(text or "")
        run.bold = bold
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        return run

    def _content_width_cm(section):
        return section.page_width.cm - section.left_margin.cm - section.right_margin.cm

    def _h1(text_value):
        paragraph = doc.add_paragraph(text_value)
        paragraph.style = doc.styles['Heading 1']
        return paragraph

    def _new_landscape_section(title):
        section = doc.add_section(WD_SECTION_START.NEW_PAGE)
        _set_orientation(section, WD_ORIENT.LANDSCAPE)
        _h1(title)
        return section

    def _ensure_portrait_section(title):
        section = doc.add_section(WD_SECTION_START.NEW_PAGE)
        _set_orientation(section, WD_ORIENT.PORTRAIT)
        _h1(title)
        return section

    def _gantt_png_buf():
        import matplotlib.pyplot as plt

        df = _workplan_df()
        if df.empty:
            return None

        height = min(1.2 + 0.35 * len(df), 12)
        fig, ax = plt.subplots(figsize=(11, height))
        _draw_gantt(ax, df, show_today=False)
        fig.subplots_adjust(left=0.30, right=0.995, top=0.96, bottom=0.36)
        buffer = BytesIO()
        fig.savefig(buffer, format="png", dpi=220, bbox_inches="tight")
        plt.close(fig)
        buffer.seek(0)
        return buffer

    first_section = doc.sections[0]
    _set_orientation(first_section, WD_ORIENT.PORTRAIT)

    _h1("Project Overview")

    id_info = st.session_state.get("id_info", {}) or {}
    budget_total = sum(float(r.get("total_usd") or 0.0) for r in st.session_state.get("budget", []))
    outputs_count = len(st.session_state.get("outputs", []))
    kpis_count = len(st.session_state.get("kpis", []))
    activities_count = len(st.session_state.get("workplan", []))

    overview_rows = [
        (LABELS["title"], id_info.get("title", "")),
        (LABELS["pi_name"], id_info.get("pi_name", "")),
        (LABELS["pi_email"], id_info.get("pi_email", "")),
        (LABELS["implementing_partners"], id_info.get("implementing_partners", "")),
        (LABELS["supporting_partners"], id_info.get("supporting_partners", "")),
        (LABELS["start_date"], fmt_dd_mmm_yyyy(id_info.get("start_date"))),
        (LABELS["end_date"], fmt_dd_mmm_yyyy(id_info.get("end_date"))),
        (LABELS["location"], id_info.get("location", "")),
        (LABELS["contact_name"], id_info.get("contact_name", "")),
        (LABELS["contact_email"], id_info.get("contact_email", "")),
        (LABELS["contact_phone"], id_info.get("contact_phone", "")),
        (LABELS["total_funding"], f"USD {budget_total:,.2f}"),
        (LABELS["outputs_count"], str(outputs_count)),
        (LABELS["kpis_count"], str(kpis_count)),
        (LABELS["activities_count"], str(activities_count)),
    ]

    overview_table = doc.add_table(rows=len(overview_rows), cols=2)
    overview_table.style = "Table Grid"
    overview_table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for idx, width in enumerate((Cm(6.0), Cm(11.0))):
        for row in overview_table.rows:
            row.cells[idx].width = width
        overview_table.columns[idx].width = width

    for row_idx, (label, value) in enumerate(overview_rows):
        # Left column: shaded + white text (same PRIMARY_SHADE used elsewhere)
        _set_cell_text(overview_table.cell(row_idx, 0), str(label), bold=True, white=True)
        _shade(overview_table.cell(row_idx, 0), PRIMARY_SHADE)

        # Right column: normal
        _set_cell_text(overview_table.cell(row_idx, 1), _s(value))

    doc.add_page_break()
    _h1("Logframe")

    goal = st.session_state.impacts[0] if st.session_state.get("impacts") else {}
    goal_text = goal.get("name", "")
    goal_assumptions_text = _assumptions_doc_text(goal.get("assumptions")) if goal else ""
    outcome_text = (st.session_state.outcomes[0]["name"] if st.session_state.get("outcomes") else "")
    out_nums, kpi_nums = compute_numbers()

    def _sort_by_num(label):
        if not label:
            return (9999,)
        try:
            return tuple(int(x) for x in str(label).split("."))
        except Exception:
            return (9999,)

    if goal_text or goal_assumptions_text:
        tbl_goal = doc.add_table(rows=1, cols=2)
        tbl_goal.style = "Table Grid"
        tbl_goal.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr = tbl_goal.rows[0]
        for idx, lab in enumerate(("Goal", "Key Assumptions")):
            _set_cell_text(hdr.cells[idx], lab, bold=True, white=True)
            _shade(hdr.cells[idx], PRIMARY_SHADE)
        _repeat_header(hdr)

        col_widths_goal = (Cm(12.0), Cm(12.0))
        for idx, width in enumerate(col_widths_goal):
            hdr.cells[idx].width = width
            tbl_goal.columns[idx].width = width

        body_row = tbl_goal.add_row()
        for idx, width in enumerate(col_widths_goal):
            body_row.cells[idx].width = width
        _set_cell_text(body_row.cells[0], goal_text or "—")
        _set_cell_text(body_row.cells[1], goal_assumptions_text or "—")
        doc.add_paragraph("")

    outcome_kpis = [k for k in st.session_state.get("kpis", []) if k.get("parent_level") == "Outcome"]
    if outcome_kpis:
        outcome_kpi = outcome_kpis[0]
        tbl_outcome = doc.add_table(rows=1, cols=4)
        tbl_outcome.style = "Table Grid"
        tbl_outcome.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr = tbl_outcome.rows[0]
        for idx, lab in enumerate(("Outcome", "KPI", "Means of Verification", "Key Assumptions")):
            _set_cell_text(hdr.cells[idx], lab, bold=True, white=True)
            _shade(hdr.cells[idx], PRIMARY_SHADE)
        _repeat_header(hdr)

        col_widths_outcome = (Cm(6.0), Cm(9.0), Cm(6.0), Cm(6.0))
        for idx, width in enumerate(col_widths_outcome):
            for row in tbl_outcome.rows:
                row.cells[idx].width = width
            tbl_outcome.columns[idx].width = width

        row = tbl_outcome.add_row()
        for idx, width in enumerate(col_widths_outcome):
            row.cells[idx].width = width

        _set_cell_text(row.cells[0], outcome_text or "")

        k_cell = row.cells[1]
        k_cell.text = ""
        para = k_cell.paragraphs[0]
        _add_run(para, f"Outcome KPI — {outcome_kpi.get('name','')}")
        para.add_run("\n")
        baseline = (outcome_kpi.get("baseline", "") or "").strip()
        if baseline:
            _add_run(para, "Baseline: ", True)
            _add_run(para, baseline)
            para.add_run("\n")
        target = (outcome_kpi.get("target", "") or "").strip()
        if target:
            _add_run(para, "Target: ", True)
            _add_run(para, target)
            para.add_run("\n")

        _set_cell_text(row.cells[2], (outcome_kpi.get("mov") or "").strip() or "—")
        _set_cell_text(row.cells[3], "—")
        doc.add_paragraph("")

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = tbl.rows[0]
    for idx, lab in enumerate(("Output", "KPI", "Means of Verification", "Key Assumptions")):
        _set_cell_text(hdr.cells[idx], lab, bold=True, white=True)
        _shade(hdr.cells[idx], PRIMARY_SHADE)
    _repeat_header(hdr)

    col_widths_main = (Cm(6.0), Cm(9.0), Cm(6.0), Cm(6.0))
    for idx, width in enumerate(col_widths_main):
        for row in tbl.rows:
            row.cells[idx].width = width
        tbl.columns[idx].width = width

    outputs = sorted(st.session_state.get("outputs", []), key=lambda o: _sort_by_num(out_nums.get(o.get("id"))))
    for output in outputs:
        out_id = output.get("id")
        out_label = out_nums.get(out_id)
        out_name = output.get("name", "")
        assumptions_txt = _assumptions_doc_text(output.get("assumptions"))
        kpi_rows = [k for k in st.session_state.get("kpis", []) if k.get("parent_id") == out_id]
        kpi_rows = sorted(kpi_rows, key=lambda k: _sort_by_num(kpi_nums.get(k.get("id"))))

        if not kpi_rows:
            row = tbl.add_row()
            for idx, width in enumerate(col_widths_main):
                row.cells[idx].width = width
            label = f"Output {out_label}" if out_label else "Output"
            if out_name:
                label = f"{label} — {out_name}"
            _set_cell_text(row.cells[0], label)
            _set_cell_text(row.cells[1], "—")
            _set_cell_text(row.cells[2], "—")
            _set_cell_text(row.cells[3], assumptions_txt or "—")
            continue

        first_row_idx = len(tbl.rows)
        for kpi in kpi_rows:
            row = tbl.add_row()
            for idx, width in enumerate(col_widths_main):
                row.cells[idx].width = width

            label = f"KPI {kpi_nums.get(kpi.get('id'), '')}".strip()
            name = kpi.get("name", "")
            k_cell = row.cells[1]
            k_cell.text = ""
            para = k_cell.paragraphs[0]
            if label and name:
                _add_run(para, f"{label} — {name}")
            else:
                _add_run(para, name or label)
            para.add_run("\n")

            baseline = (kpi.get("baseline") or "").strip()
            if baseline:
                _add_run(para, "Baseline: ", True)
                _add_run(para, baseline)
                para.add_run("\n")
            target = (kpi.get("target") or "").strip()
            if target:
                _add_run(para, "Target: ", True)
                _add_run(para, target)
                para.add_run("\n")

            mov = (kpi.get("mov") or "").strip()
            _set_cell_text(row.cells[2], mov or "—")
            _set_cell_text(row.cells[3], _assumptions_doc_text(kpi.get("assumptions")) or "—")

        last_row_idx = len(tbl.rows) - 1
        if last_row_idx >= first_row_idx:
            merged_label = tbl.cell(first_row_idx, 0).merge(tbl.cell(last_row_idx, 0))
            label = f"Output {out_label}" if out_label else "Output"
            if out_name:
                label = f"{label} — {out_name}"
            _set_cell_text(merged_label, label)
            merged_assumption = tbl.cell(first_row_idx, 3).merge(tbl.cell(last_row_idx, 3))
            _set_cell_text(merged_assumption, assumptions_txt or "—")

    doc.add_paragraph("")

    workplan_section = _new_landscape_section("Workplan")
    df_wp = _workplan_df()
    if df_wp.empty:
        doc.add_paragraph("No activities defined yet.")
    else:
        df_wp = df_wp.copy()
        df_wp["Start"] = pd.to_datetime(df_wp["Start"], errors="coerce").dt.strftime("%d/%b/%Y")
        df_wp["End"] = pd.to_datetime(df_wp["End"], errors="coerce").dt.strftime("%d/%b/%Y")
        df_wp["Start"] = df_wp["Start"].fillna("")
        df_wp["End"] = df_wp["End"].fillna("")

        col_widths_act = (Cm(5.5), Cm(6.4), Cm(4.6), Cm(3.0), Cm(3.0))
        t_act = doc.add_table(rows=1, cols=5)
        t_act.style = "Table Grid"
        t_act.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr = t_act.rows[0]
        for idx, lab in enumerate(("Output", "Activity", "Owner", "Start date", "End date")):
            _set_cell_text(hdr.cells[idx], lab, bold=True, white=True)
            _shade(hdr.cells[idx], PRIMARY_SHADE)
        _repeat_header(hdr)
        for idx, width in enumerate(col_widths_act):
            for row in t_act.rows:
                row.cells[idx].width = width
            t_act.columns[idx].width = width

        out_nums_map, _ = compute_numbers()
        name_to_num = {}
        for output in st.session_state.get("outputs", []):
            name_to_num[output.get("name")] = out_nums_map.get(output.get("id"))

        grouped = df_wp.groupby("Output")
        for out_name, subset in grouped:
            subset = subset.sort_values("ActivityLabel")
            start_row = len(t_act.rows)
            for _, activity_row in subset.iterrows():
                row = t_act.add_row()
                for idx, width in enumerate(col_widths_act):
                    row.cells[idx].width = width
                _set_cell_text(row.cells[1], str(activity_row["ActivityLabel"]))
                owner_val = activity_row.get("Owner", "")
                if isinstance(owner_val, float):
                    try:
                        from math import isnan
                        if isnan(owner_val):
                            owner_val = ""
                    except Exception:
                        owner_val = ""
                owner_text = str(owner_val).strip()
                if not owner_text or owner_text.lower() == "nan":
                    owner_text = "—"
                _set_cell_text(row.cells[2], owner_text)
                _set_cell_text(row.cells[3], activity_row["Start"] or "")
                _set_cell_text(row.cells[4], activity_row["End"] or "")
            end_row = len(t_act.rows) - 1
            if end_row >= start_row:
                merged = t_act.cell(start_row, 0).merge(t_act.cell(end_row, 0))
                num = name_to_num.get(out_name, "")
                if num and out_name:
                    label = f"Output {num} — {out_name}"
                elif num:
                    label = f"Output {num}"
                else:
                    label = out_name or "Output"
                _set_cell_text(merged, label)

    doc.add_paragraph("")

    width_cm = _content_width_cm(workplan_section)
    gantt_source = gantt_image_path or None
    if gantt_source is None:
        try:
            gantt_source = _gantt_png_buf()
        except Exception:
            gantt_source = None

    if gantt_source:
        try:
            heading_style = doc.styles['Heading 2']
        except KeyError:
            heading_style = doc.styles['Heading 1']
        doc.add_paragraph("Workplan Gantt", style=heading_style)
        try:
            from pathlib import Path as _Path
            pic_source = str(gantt_source) if isinstance(gantt_source, (str, _Path)) else gantt_source
            doc.add_picture(pic_source, width=Cm(max(1.0, width_cm - 0.5)))
        except Exception:
            pass

    doc.add_paragraph("")

    _new_landscape_section("Budget")
    bt = doc.add_table(rows=1, cols=8)
    bt.style = "Table Grid"
    bt.alignment = WD_TABLE_ALIGNMENT.LEFT
    bh = bt.rows[0]
    headers = ("Activity", "Line Item", "Category", "Sub Category", "Unit", "Unit Cost (USD)", "Quantity", "Total (USD)")
    for idx, lab in enumerate(headers):
        _set_cell_text(bh.cells[idx], lab, bold=True, white=True)
        _shade(bh.cells[idx], PRIMARY_SHADE)
    _repeat_header(bh)

    colw = (Cm(5.2), Cm(5.5), Cm(3.8), Cm(4.2), Cm(2.8), Cm(3.2), Cm(2.8), Cm(3.4))
    for idx, width in enumerate(colw):
        for row in bt.rows:
            row.cells[idx].width = width
        bt.columns[idx].width = width

    _, _, act_nums = compute_numbers(include_activities=True)
    act_lookup = activity_label_map(act_nums)

    total_budget = 0.0
    for budget_row in st.session_state.get("budget", []):
        row = bt.add_row()
        for idx, width in enumerate(colw):
            row.cells[idx].width = width
        _set_cell_text(row.cells[0], act_lookup.get(budget_row.get("activity_id"), ""))
        _set_cell_text(row.cells[1], budget_row.get("item", ""))
        _set_cell_text(row.cells[2], budget_row.get("category", ""))
        _set_cell_text(row.cells[3], budget_row.get("subcategory", ""))
        _set_cell_text(row.cells[4], budget_row.get("unit", ""))

        unit_cost = float(budget_row.get("unit_cost") or 0.0)
        quantity = float(budget_row.get("qty") or 0.0)
        total = float(budget_row.get("total_usd") or (unit_cost * quantity) or 0.0)
        total_budget += total

        _set_cell_text(row.cells[5], f"{unit_cost:,.2f}")
        _set_cell_text(row.cells[6], f"{quantity:,.2f}")
        _set_cell_text(row.cells[7], f"{total:,.2f}")

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run(f"Total budget: USD {total_budget:,.2f}")
    run.bold = True

    note = doc.add_paragraph()
    note_run = note.add_run(footer_note)
    note_run.italic = True

    _ensure_portrait_section("Disbursement Schedule")
    from datetime import date as _date

    dsp_src = list(st.session_state.get("disbursement", []))
    out_nums_map, _ = compute_numbers()
    id_to_output = {o["id"]: (o.get("name") or "Output") for o in st.session_state.outputs}

    def _out_label(output_id):
        return f"Output {out_nums_map.get(output_id, '')} — {id_to_output.get(output_id, '(unassigned)')}".strip(" —")

    def _out_num_val(output_id):
        label = out_nums_map.get(output_id, "")
        try:
            return int(label)
        except Exception:
            return 10 ** 9

    dsp_rows = sorted(
        dsp_src,
        key=lambda row: (
            row.get("anticipated_date") or _date(2100, 1, 1),
            _out_num_val(row.get("output_id")),
            (row.get("kpi_name") or row.get("deliverable") or "").strip(),
        ),
    )

    t_disb = doc.add_table(rows=1, cols=5)
    t_disb.style = "Table Grid"
    t_disb.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t_disb.rows[0]
    for idx, lab in enumerate(("Output", "KPI", "Anticipated deliverable date", "Deliverable", "Amount (USD)")):
        _set_cell_text(hdr.cells[idx], lab, bold=True, white=True)
        _shade(hdr.cells[idx], PRIMARY_SHADE)
    _repeat_header(hdr)

    disb_widths = (Cm(3.5), Cm(4.0), Cm(3.5), Cm(7.0), Cm(3.0))
    for idx, width in enumerate(disb_widths):
        for row in t_disb.rows:
            row.cells[idx].width = width
        t_disb.columns[idx].width = width

    if dsp_rows:
        for d_row in dsp_rows:
            row = t_disb.add_row()
            for idx, width in enumerate(disb_widths):
                row.cells[idx].width = width
            _set_cell_text(row.cells[0], _out_label(d_row.get("output_id")))
            _set_cell_text(row.cells[1], d_row.get("kpi_name") or "")
            _set_cell_text(row.cells[2], fmt_dd_mmm_yyyy(d_row.get("anticipated_date")) or "")
            _set_cell_text(row.cells[3], d_row.get("deliverable") or "")
            _set_cell_text(row.cells[4], f"{float(d_row.get('amount_usd') or 0.0):,.2f}")
    else:
        row = t_disb.add_row()
        for idx, width in enumerate(disb_widths):
            row.cells[idx].width = width
        _set_cell_text(row.cells[0], "")
        _set_cell_text(row.cells[1], "")
        _set_cell_text(row.cells[2], "")
        _set_cell_text(row.cells[3], "No disbursements defined.")
        _set_cell_text(row.cells[4], "")

    assumption_rows = []
    if goal_assumptions_text:
        assumption_rows.append(("Goal", goal_assumptions_text))
    for outcome in st.session_state.get("outcomes", []):
        text_value = _assumptions_doc_text(outcome.get("assumptions"))
        if text_value:
            assumption_rows.append((f"Outcome — {outcome.get('name','')}", text_value))
    for output in st.session_state.get("outputs", []):
        text_value = _assumptions_doc_text(output.get("assumptions"))
        if text_value:
            label = f"Output {out_nums.get(output.get('id'), '')} — {output.get('name','')}".strip(" —")
            assumption_rows.append((label or "Output", text_value))

    if assumption_rows:
        doc.add_paragraph("")
        try:
            heading_style = doc.styles['Heading 2']
        except KeyError:
            heading_style = doc.styles['Heading 1']
        doc.add_paragraph("Assumptions", style=heading_style)
        tbl_ass = doc.add_table(rows=len(assumption_rows), cols=2)
        tbl_ass.style = "Table Grid"
        tbl_ass.alignment = WD_TABLE_ALIGNMENT.LEFT
        for idx, width in enumerate((Cm(5.0), Cm(11.0))):
            for row in tbl_ass.rows:
                row.cells[idx].width = width
            tbl_ass.columns[idx].width = width
        for row_idx, (label, text_value) in enumerate(assumption_rows):
            _set_cell_text(tbl_ass.cell(row_idx, 0), label, bold=True)
            _set_cell_text(tbl_ass.cell(row_idx, 1), text_value)

    note = doc.add_paragraph()
    note_run = note.add_run(footer_note)
    note_run.italic = True

    _force_calibri_everywhere(doc)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def view_activity_readonly(a, label, id_to_output, id_to_kpi):
    out_name = id_to_output.get(a.get("output_id"), "(unassigned)")
    kpis_txt = ", ".join(id_to_kpi.get(kid, "") for kid in (a.get("kpi_ids") or [])) or "—"
    sd = fmt_dd_mmm_yyyy(a.get("start")) or "—"
    ed = fmt_dd_mmm_yyyy(a.get("end"))   or "—"

    body = (
        f"<div class='lf-activity-title'>Activity {escape(label)} — {escape(a.get('name',''))}</div>"
        f"<div class='lf-line'><b>Output:</b> {escape(out_name)}</div>"
        f"<div class='lf-line'><b>Owner:</b> {escape(a.get('owner','') or '—')}</div>"
        f"<div class='lf-line'><b>Start date:</b> {sd} &nbsp;&nbsp;•&nbsp;&nbsp; <b>End date:</b> {ed}</div>"
        f"<div class='lf-line'><b>Linked KPIs:</b> {escape(kpis_txt)}</div>"
    )

    return view_logframe_element(body, kind="activity")


@contextmanager
def lf_card_container(*extra_classes: str):
    """Render a div with the shared lf-card styling and always close it."""
    classes = " ".join(["lf-card", *extra_classes]) if extra_classes else "lf-card"
    st.markdown(f"<div class='{classes}'>", unsafe_allow_html=True)
    try:
        yield
    finally:
        st.markdown("</div>", unsafe_allow_html=True)

def view_budget_item_card(row, id_to_output) -> str:
    """
    row: [OutputID, Item, Category, Unit, Qty, Unit Cost, Currency, Total]
    id_to_output: {output_id -> output name}
    """
    out_id, item, cat, unit, qty, uc, cur, tot, blid = _budget_unpack(row)
    out_name = id_to_output.get(out_id, "(unassigned)")
    # rows (label/value)
    lines = [
        ("Output", out_name),
        ("Item", item or "—"),
        ("Category", cat or "—"),
        ("Unit", unit or "—"),
        ("Qty", str(qty or 0)),
        ("Unit Cost", fmt_money(uc or 0)),
        ("Currency", cur or "—"),
        ("Total", fmt_money(tot or 0)),
    ]
    body = "".join(
        f"<div class='lf-line'><b>{escape(lbl)}:</b> {escape(val) if lbl not in ('Unit Cost','Total') else val}</div>"
        for (lbl, val) in lines
    )
    # reuse the blue style
    return f"<div class='lf-card lf-card--budget'>{body}</div>"

# ---------------- CSS for cards ----------------
def inject_logframe_css():
    st.markdown("""
    <style>
/* Base card */
.lf-card{
  margin: 14px 0 16px;
  padding: 14px 16px;
  border-radius: 12px;
  box-shadow: 0 1px 2px rgba(0,0,0,.03);
  position: relative;               /* required for the left accent bar */
}
.lf-goal-header{ margin: 0; font-weight: 700; }

/* ========= GOAL (blue) ========= */
.lf-card--goal{
  background: #E6F0FB;
  border: 1px solid #7BA4D9;
}
.lf-card--goal::before{
  content:"";
  position:absolute; top:0; left:0; bottom:0; width:8px;
  background:#4F7BB5;
  border-top-left-radius:12px;
  border-bottom-left-radius:12px;
}

/* ========= OUTPUT (green) ========= */
.lf-card--output{
  background: #E9F3E1;              /* pastel green fill */
  border: 1px solid #91A76A;        /* brand green border (#91A76A) */
}
.lf-card--output::before{
  content:"";
  position:absolute; top:0; left:0; bottom:0; width:8px;
  background:#6E8C4B;               /* darker green accent bar */
  border-top-left-radius:12px;
  border-bottom-left-radius:12px;
}

/* ========= KPI (orange) ========= */
.lf-card--kpi{
  background: #FFEBD6;          /* soft orange */
  border: 1px solid #F2B277;    /* warm orange border */
  border-radius: 12px;
  padding: 12px 16px;
  margin: 10px 0 12px;
  position: relative;
  box-shadow: 0 1px 2px rgba(0,0,0,.03);
}
.lf-card--kpi::before{
  content:"";
  position:absolute; top:0; left:0; bottom:0; width:6px;
  background:#DD7A1A;          /* stronger orange accent */
  border-top-left-radius:12px;
  border-bottom-left-radius:12px;
}

/* ========= Activity (yellow) ========= */
.lf-card--activity{
  background: #FFFBEA;          /* soft yellow */
  border: 1px solid #F6D58E;    /* warm yellow border */
  border-radius: 12px;
  padding: 12px 16px;
  margin: 10px 0 12px;
  position: relative;
  box-shadow: 0 1px 2px rgba(0,0,0,.03);
}
.lf-card--activity::before{
  content:"";
  position:absolute; top:0; left:0; bottom:0; width:6px;
  background:#F59E0B;          /* amber accent */
  border-top-left-radius:12px;
  border-bottom-left-radius:12px;
}

/* Headings & text bits */
.lf-out-header{ margin: 0; font-weight: 700; }
.lf-kpi-title{ font-weight: 600; margin-bottom: 6px; }

/* Assumptions list styling */
.lf-ass-heading{
  font-weight: 600;
  margin: 6px 0 4px;                 /* tighter heading spacing */
}

.lf-ass-list{
  margin: 4px 0 0 1.15rem;           /* small left indent, no extra top gap */
  padding: 0;
  list-style: disc outside;
}

.lf-ass-list li{
  margin: 0;                          /* remove extra gaps between items */
  padding: 0;
  font-style: italic;                 /* italics */
  font-size: 0.92rem;                 /* smaller text */
  line-height: 1.25;                  /* single/compact spacing */
}

/* Chips - Check if linked to payment*/
.chip {
  display:inline-block; padding: 2px 8px; border-radius: 999px;
  background:#eef2ff; font-size: .85rem; border:1px solid #e2e8ff;
}
.chip.green { background:#e6f9ee; border-color:#c6f0d8; color:#046c4e; }

/* Dot bullet in Output header */
.dot { font-size: 1.05rem; margin-right: .4rem; }

/* (optional) generic line styling used in KPI details */
.lf-line { margin: 2px 0; font-size: 0.95rem; color: #444; }

/* ========= BUDGET (blue) ========= */
.lf-card--budget{
  background: #E8F0FE;           /* light blue */
  border: 1px solid #90B4FE;
  border-radius: 12px;
  padding: 12px 16px;
  margin: 10px 0 12px;
  box-shadow: 0 1px 2px rgba(0,0,0,.03);
}
.lf-budget-total{ font-weight:700; margin-top:10px; }

.lf-card--budget .budget-table td:nth-child(1){
  max-width: 520px;              /* tune as needed */
  white-space: nowrap;
  overflow: hidden;
  text-overflow: ellipsis;
}   

/* --- Compact budget row (blue) --- */
.lf-budget-row{
  background:#E8F0FE;
  border:1px solid #90B4FE;
  border-radius:12px;
  padding:10px 12px;
  margin:8px 0;
  display:flex;
  align-items:center;
  gap:10px;
  box-shadow:0 1px 2px rgba(0,0,0,.03);
}
.lf-budget-cells{
  display:flex;
  gap:14px;
  flex-wrap:wrap;
  align-items:baseline;
  width:100%;
  font-size:.95rem;
  color:#1f2937;
}
.lf-budget-cell{
  white-space:nowrap;
}
.lf-budget-money{
  font-variant-numeric: tabular-nums;
  text-align:right;
}
.lf-budget-chip{
  background:#eef2ff;
  border:1px solid #dbe2ff;
  border-radius:999px;
  padding:2px 8px;
  display:inline-block;
}
.lf-subtotal{ font-weight:700; margin:6px 0 12px; text-align:right; }
.lf-grandtotal{ font-weight:800; margin-top:12px; text-align:right; }

/* Make the Streamlit-native header logo larger */
header [data-testid="stLogo"] img {
  height: 80px;      /* change this */
  width: auto;
}
header [data-testid="stLogo"] {
  margin-left: 8px;  /* optional: shift slightly from edge */
  padding-top: 6px;  /* optional */
  padding-bottom: 6px;
}

/* Budget tab formatting */
<style>
/* Column headers bolder */
div[data-testid="column"] label p { font-weight: 600; }

/* Right align helper (we'll use it for Totals) */
.right { text-align: right; }

/* Wrap long labels/text inside selectboxes and their option list */
.stSelectbox div[role="combobox"] span, 
.stSelectbox [role="listbox"] div {
  white-space: normal !important;
  line-height: 1.25 !important;
}

/* Make the visible select value container wrap (Streamlit uses react-select) */
.stSelectbox div[aria-haspopup="listbox"] > div {
  white-space: normal !important;
}

/* Reduce excess vertical gaps between rows a bit */
.block-container .stHorizontalBlock { margin-bottom: 0.35rem; }

/* Total field look: plain text, bold, no input-like border */
.budget-total {
  font-weight: 700; 
  display: inline-block; 
  padding: .4rem .6rem; 
  background: #fff; 
  border-radius: 6px;
  border: 1px solid rgba(0,0,0,.08);

    </style>
    """, unsafe_allow_html=True)

CATEGORY_TREE = {
    "Personnel": [
        "Salaries & Wages",
        "Temporary / Short-Term Staff"
    ],

    "Professional Services": [
        "Consultancy / Advisory Services",
        "Research & Data Collection",
        "Monitoring, Evaluation & Learning",
        "IT & Systems Development",
        "Legal, Audit & Compliance",
        "Translation & Interpretation"
    ],

    "Travel & Transportation": [
        "Airfare",
        "Travel insurance",
        "Visa fees",
        "Accommodation & Lodging",
        "Ground Transportation",
        "Per Diem / Allowances"
    ],
    
    "Events, Workshops, Training & Meetings": [
        "Venue Hire",
        "Facilitation & Trainer Fees",
        "Audio/Visual & Interpretation (on-site)",
        "Training / Workshop Materials",
        "Catering & Refreshments",
        "Event Packages (e.g. entertainment, gifts, souvernirs, awards, etc.)",
        "Per diem"
    ],
    
    "Equipment & Technology": [
        "ICT Equipment (computers, phones, tablets)",
        "Software & Licenses",
        "Power & Connectivity (chargers, solar, batteries)",
        "Lab / Medical Equipment",
        "Hosting & Maintenance"
    ],

    "Supplies & Materials": [
        "Training Materials",
        "Printing & Dissemination",
        "Stationery & Office Supplies",
        "Safety / Field Gear",
        "Lab / Medical Consumables",
        "Other Consumables"
    ],

    "Field Operations": [
        "Transport & Vehicle Operations",
        "Warehousing & Storage",
        "Permits & Local Fees"
    ],

    "Communications & Advocacy": [
        "Printing & Dissemination",
        "Media Campaigns",
        "Design & Branding",
        "Publications & Layout",
        "Stakeholder Engagement / Public Relations"
    ],
    
    "Administrative / Direct Operating Costs": [
        "Office Rent & Utilities",
        "Project Communications (internet/phone)",
        "Insurance",
        "General Admin & Shared Services"
    ],

    "Indirect Costs / Overheads": [
        "Institutional Overhead (as per GLIDE policy)"
    ]
}

# Sub Category–driven Unit suggestions (bind only to Sub Category)
SUBCATEGORY_UNIT_SUGGESTIONS = {
    # Personnel
    "Salaries & Wages": ["per hour", "per day", "per month", "per FTE-month"],
    "Temporary / Short-Term Staff": ["per day", "per hour", "per contract"],
    
    # Events, Workshops, Training & Meetings
    "Venue Hire":                              ["per day", "per event"],
    "Facilitation & Trainer Fees":             ["per day", "per event"],
    "Audio/Visual & Interpretation (on-site)": ["per day", "per event"],
    "Training / Workshop Materials":           ["per participant", "per event"],
    "Catering & Refreshments":                 ["per participant", "per event day", "per event"],
    "Event Packages (e.g. entertainment, gifts, souvernirs, awards, etc.)": ["per event", "per person"],
    "Per Diem / Allowances":                   ["per participant per day"],

    # Travel & Transportation
    "Airfare":                                 ["per individual"],
    "Travel insurance":                        ["per individual"],
    "Visa fees":                               ["per individual"],
    "Accommodation & Lodging":                 ["per night"],
    "Ground Transportation":                   ["per person", "per trip", "per km"],
    "Per Diem / Allowances":                   ["per participant per day"],

    # Equipment & Technology
    "ICT Equipment (e.g., computers, phones, tablets)": ["per device"],
    "Software & Licenses":                     ["per license", "per month", "per year"],
    "Power & Connectivity (e.g., chargers, solar, batteries)": ["per item", "per set"],
    "Lab / Medical Equipment":                 ["per device"],
    "Hosting & Maintenance":                   ["per month", "per year"],

    # Supplies & Materials
    "Training Materials":                      ["per participant", "per set", "per item"],
    "Printing & Dissemination":                ["per copy", "per batch"],
    "Stationery & Office Supplies":            ["per item", "per set", "per batch"],
    "Safety / Field Gear":                     ["per set", "per item", "per person"],
    "Lab / Medical Consumables":               ["per item", "per batch"],
    "Other Consumables":                       ["per item", "per pack", "per batch"],

    # Field Operations
    "Transport & Vehicle Operations":          ["per vehicle-day", "per trip"],
    "Warehousing & Storage":                   ["per m²-month", "per pallet-month"],
    "Permits & Local Fees":                    ["per item", "lump sum"],

    # Communications & Advocacy
    "Printing & Dissemination":                ["per copy", "per 1,000 copies", "per batch"],
    "Media Campaigns":                         ["per campaign", "per campaign day", "per 30-second spot"],
    "Design & Branding":                       ["per design"],
    "Publications & Layout":                   ["per report", "per page"],
    "Stakeholder Engagement / Public Relations": ["per event", "per campaign"],

    # Professional Services
    "Consultancy / Advisory Services":         ["per day", "per contract"],
    "Research & Data Collection":              ["per day", "per survey", "per dataset"],
    "Monitoring, Evaluation & Learning":       ["per day", "per evaluation", "per study"],
    "IT & Systems Development":                ["per day", "per contract"],
    "Legal, Audit & Compliance":               ["per day", "per contract"],
    "Translation & Interpretation":            ["per day", "per translation page", "per 1,000 words"],

    # Administrative / Direct Operating Costs
    "Office Rent & Utilities":                 ["per month"],
    "Project Communications (internet/phone)": ["per month"],
    "Insurance":                               ["per month", "per year"],
    "General Admin & Shared Services":         ["per month"],

    # Indirect Costs / Overheads
    "Institutional Overhead (as per GLIDE policy)": ["lump sum"],
}

DEFAULT_UNIT_CHOICES = ["per item", "per unit"]
UNIT_PLACEHOLDER = "Select unit"
CUSTOM_UNIT_OPTION = "Custom…"

SUBCATEGORY_UNITS = {
    (category, sub): SUBCATEGORY_UNIT_SUGGESTIONS.get(sub, [])
    for category, subs in CATEGORY_TREE.items()
    for sub in subs
}


def unit_choices_for(category: str | None, subcategory: str) -> tuple[list[str], bool]:
    """Return (unit options, has_defined_mapping) for the given category/subcategory."""
    defined = SUBCATEGORY_UNITS.get((category, subcategory), []) if category is not None else []
    has_defined = bool(defined)
    base = defined if has_defined else SUBCATEGORY_UNIT_SUGGESTIONS.get(subcategory, [])
    if not base:
        base = DEFAULT_UNIT_CHOICES
    seen, ordered = set(), []
    for u in base:
        if u and u not in seen:
            seen.add(u)
            ordered.append(u)
    return ordered, has_defined


def unit_options_for_sub(subcategory: str, category: str | None = None, *, include_custom: bool = True) -> list[str]:
    units, _ = unit_choices_for(category, subcategory)
    ordered = list(units)
    if include_custom and CUSTOM_UNIT_OPTION not in ordered:
        ordered.append(CUSTOM_UNIT_OPTION)
    return ordered

def subcategories_for(category: str):
    return CATEGORY_TREE.get(category, [])

def _budget_unpack(row):
    """Compat helper -> returns a 9-tuple for old/new rows."""
    return (
        row.get("output_id"),
        row.get("item"),
        row.get("category"),
        row.get("unit"),
        row.get("qty"),
        row.get("unit_cost"),
        row.get("currency", "USD"),
        row.get("total_usd"),
        row.get("id"),
    )
# ---------------- Tabs ----------------
tabs = st.tabs([
    "📘 Welcome & Instructions",
    "🪪 Project Overview",
    "🧱 Logframe",
    "🗂️ Workplan",
    "💵 Budget",
    "💸 Disbursement Schedule",
    "📤 Export"
])

# ===== TAB 1: Welcome & Instructions =====
tabs[0].markdown(
    """
# 📝 Welcome to the Falcon Awards Project Portal

Please complete each section of your project:

1. **Project Overview** – Fill project & contacts.
2. **Logframe** – Add **Goal**, **Outcome**, **Outputs**, then **KPIs**.
3. **Workplan** – Add **Activities** linked to Outputs/KPIs with dates.
4. **Budget** – Add **Budget lines** linked to Activities.
5. **Disbursement Schedule** – Add **KPI-Linked Milestones** with dates and amounts.
6. **Export** – Use the export tools as follows:  
   - **Create a backup file (Excel):** You can do this at any stage to save your progress and continue working later.  
   - **Generate a Project Design Document (Word):** Create a formatted project document to review how your application looks in report form.  
   - **Generate a Project Package (ZIP):** At the final stage you will create a password-protected ZIP file for submission to GLIDE. *This feature is planned but not yet available in the portal.*  

### Definitions
- **Goal**: The long-term vision (impact).
- **Outcome**: The specific change expected from the project.
- **Output**: Tangible products/services delivered by the project.
- **Key Performance Indicator (KPI)**: Quantifiable metric to judge performance (with baseline, target, dates, MoV).
- **Activity**: Tasks that produce Outputs (scheduled in the Workplan).
- **Budget line**: A costed item linked to an activity (category, unit, quantity, unit cost). Together, the budget lines under an activity represent the resources required to implement it.
- **Assumptions**: External conditions necessary for success.
- **Means of Verification (MoV)**: Where/how the KPI will be measured.
- **Payment-linked indicator**: KPI that triggers funding release when achieved (optional).

"""
)
# **Contact us**: Anderson E. Stanciole | astanciole@glideae.org

# --- Resume from Excel ---
uploaded_file = tabs[0].file_uploader("Resume Previous Submission (Excel)", type="xlsx")
if uploaded_file is not None:
    try:
        # Build a stable signature for the uploaded file
        file_bytes = uploaded_file.getvalue()
        file_sig = hashlib.md5(file_bytes).hexdigest()

        # Only (re)load if this is a new file or changed content
        if st.session_state.get("_resume_file_sig") != file_sig:
            xls = pd.ExcelFile(BytesIO(file_bytes))

            # ---- RESET state containers
            st.session_state.impacts = []
            st.session_state.outcomes = []
            st.session_state.outputs = []
            st.session_state.kpis = []

            # (optional) clear edit flags so they won't point to old IDs
            for _f in ("edit_goal", "edit_outcome", "edit_output", "edit_kpi"):
                st.session_state[_f] = None

            # ---- Read Summary with explicit IDs ----
            summary_df = pd.read_excel(xls, sheet_name="Summary")
            summary_df.columns = [str(c).strip() for c in summary_df.columns]

            st.session_state.impacts = []
            st.session_state.outcomes = []
            st.session_state.outputs = []

            for _, row in summary_df.iterrows():
                lvl = _s(row.get("RowLevel", ""))
                _id = _s(row.get("ID")) or generate_id()
                pid = _s(row.get("ParentID")) or None
                text = _s(row.get("Text / Title"))
                ass = _s(row.get("Assumptions"))

                if lvl.lower() == "goal":
                    st.session_state.impacts.append({
                        "id": _id,
                        "level": "Goal",
                        "name": text,
                        "assumptions": ass,
                    })

                elif lvl.lower() == "outcome":
                    st.session_state.outcomes.append({
                        "id": _id,
                        "level": "Outcome",
                        "name": text,
                        "parent_id": pid,
                        "assumptions": ass,
                    })

                elif lvl.lower() == "output":
                    st.session_state.outputs.append({
                        "id": _id,
                        "level": "Output",
                        "name": strip_label_prefix(text, "Output") or "Output",
                        "parent_id": pid,
                        "assumptions": ass
                    })

            # ---- KPI Matrix (ID-based) ----
            if "KPI Matrix" in xls.sheet_names:
                kdf = pd.read_excel(xls, sheet_name="KPI Matrix")
                kdf.columns = [str(c).strip() for c in kdf.columns]

                st.session_state.kpis = []  # reset before loading

                for _, r in kdf.iterrows():
                    kid = _s(r.get("KPIID")) or generate_id()
                    plev = _s(r.get("Parent Level") or "Output")
                    pid = _s(r.get("ParentID")) or None

                    st.session_state.kpis.append({
                        "id": kid,
                        "level": "KPI",
                        "name": _s(r.get("KPI")),
                        "parent_id": pid,
                        "parent_level": "Outcome" if plev.lower() == "outcome" else "Output",
                        "baseline": _s(r.get("Baseline")),
                        "target": _s(r.get("Target")),
                        "start_date": parse_date_like(r.get("Start Date")),
                        "end_date": parse_date_like(r.get("End Date")),
                        "linked_payment": _s(r.get("Linked to Payment")).lower() in ("yes", "y", "true", "1"),
                        "mov": _s(r.get("Means of Verification")),
                    })

            # ---- Workplan (supports both "rich" and "simple" exports)
            if "Workplan" in xls.sheet_names:
                wdf = pd.read_excel(xls, sheet_name="Workplan")
                # normalize headers
                wdf.columns = [str(c).strip() for c in wdf.columns]

                def _split_csv(s):
                    return [t.strip() for t in (_s(s).split(",") if _s(s) else []) if t.strip()]

                # lookups
                outputs_by_name = {(o.get("name") or "").strip(): o["id"] for o in st.session_state.outputs}
                kpis_by_name = {(k.get("name") or "").strip(): k["id"] for k in st.session_state.kpis}
                kpi_id_set = set(kpis_by_name.values())

                # detect "rich" vs "simple" format
                rich = {"Activity ID", "Activity #", "OutputID", "Output", "Activity", "Owner", "Start", "End",
                        "Linked KPI IDs", "Linked KPIs"}.issubset(set(wdf.columns))

                st.session_state.workplan = []  # reset before loading

                if rich:
                    output_ids = {o["id"] for o in st.session_state.outputs}
                    kpi_ids = {k["id"] for k in st.session_state.kpis}


                    def _csv_ids(s):
                        vals = [t.strip() for t in (_s(s).split(",") if _s(s) else []) if t.strip()]
                        return vals


                    for _, row in wdf.iterrows():
                        act_id = _s(row.get("Activity ID")) or generate_id()
                        out_id = _s(row.get("OutputID"))
                        if not out_id:
                            out_name = _s(row.get("Output"))
                            out_id = next((o["id"] for o in st.session_state.outputs if
                                           (o.get("name") or "").strip() == out_name), None)

                        linked_kpi_ids = [kid for kid in _csv_ids(row.get("Linked KPI IDs")) if kid in kpi_ids]
                        dep_ids = _csv_ids(row.get("Dependencies"))

                        st.session_state.workplan.append({
                            "id": act_id,
                            "output_id": out_id if out_id in output_ids else None,
                            "name": _s(row.get("Activity")),
                            "owner": _s(row.get("Owner")),
                            "start": parse_date_like(row.get("Start")),
                            "end": parse_date_like(row.get("End"))
                        })

                else:
                    # SIMPLE legacy format: Activity | Owner | Start Date | End Date | Milestone
                    # No Output/KPI info in this shape; if there is exactly one Output, attach to it.
                    only_output_id = st.session_state.outputs[0]["id"] if len(st.session_state.outputs) == 1 else None
                    for _, row in wdf.iterrows():
                        st.session_state.workplan.append({
                            "id": generate_id(),
                            "output_id": only_output_id,  # None if multiple outputs; user can reassign in UI
                            "name": _s(row.get("Activity")),
                            "owner": _s(row.get("Owner")),
                            "start": parse_date_like(row.get("Start Date")),
                            "end": parse_date_like(row.get("End Date")),
                            "kpi_ids": [],
                        })

            # ---- Budget import (supports new detailed schema + legacy fallback)
            if "Budget" in xls.sheet_names:
                bdf = pd.read_excel(xls, sheet_name="Budget")
                bdf.columns = [str(c).strip() for c in bdf.columns]
                st.session_state.budget = []

                # helper: map activity label back to id
                # (recompute numbers after import of workplan if needed)
                _, _, act_nums = compute_numbers(include_activities=True)
                act_lookup = activity_label_map(act_nums)
                label_to_act_id = {v: k for k, v in act_lookup.items()}

                cols = set(bdf.columns)

                # New detailed format
                if {"Budget Line Item", "Cost Category", "Sub Category", "Unit",
                    "Unit Cost (USD)", "Quantity", "Total Cost (USD)"}.issubset(cols):

                    for _, r in bdf.iterrows():
                        line_item = _s(r.get("Budget Line Item"))
                        cat = _s(r.get("Cost Category"))
                        subcat = _s(r.get("Sub Category"))
                        unit = _s(r.get("Unit"))
                        uc = float(r.get("Unit Cost (USD)") or 0.0)
                        qty = float(r.get("Quantity") or 0.0)
                        tot = float(r.get("Total Cost (USD)") or (uc * qty) or 0.0)

                        # Try to recover activity_id from the label (export stores label)
                        act_label = _s(r.get("Linked Activity"))
                        activity_id = label_to_act_id.get(act_label) if act_label else None

                        st.session_state.budget.append({
                            "id": generate_id(),
                            "activity_id": activity_id,
                            "item": line_item,
                            "category": cat,
                            "subcategory": subcat,
                            "unit": unit,
                            "unit_cost": uc,
                            "qty": qty,
                            "currency": "USD",
                            "total_usd": tot,
                        })

                # Legacy simple format (3 columns)
                elif {"Budget item", "Description", "Total Cost (USD)"}.issubset(cols):
                    for _, r in bdf.iterrows():
                        item = _s(r.get("Budget item"))
                        desc = _s(r.get("Description"))
                        tot = float(r.get("Total Cost (USD)") or 0.0)
                        if item or desc or tot:
                            # map legacy to new structure (simple best-effort)
                            st.session_state.budget.append({
                                "id": generate_id(),
                                "activity_id": None,
                                "item": item or desc or "Budget line",
                                "category": "Professional Services",
                                "subcategory": "Consultancy / Advisory Services",
                                "unit": "lump sum",
                                "unit_cost": float(tot),
                                "qty": 1.0,
                                "currency": "USD",
                                "total_usd": float(tot),
                            })

            # ---- Disbursement Schedule import (optional sheet)
            if "Disbursement Schedule" in xls.sheet_names:
                ddf = pd.read_excel(xls, sheet_name="Disbursement Schedule")
                ddf.columns = [str(c).strip() for c in ddf.columns]

                st.session_state.disbursement = []

                k_by_id = {k["id"]: k for k in st.session_state.kpis}

                for _, row in ddf.iterrows():
                    kid = _s(row.get("KPIID"))
                    k = k_by_id.get(kid)

                    st.session_state.disbursement.append({
                        "kpi_id": kid,
                        "output_id": (k.get("parent_id") if k else None),
                        "kpi_name": (k.get("name") if k else ""),
                        "anticipated_date": parse_date_like(row.get("Anticipated deliverable date")),
                        "deliverable": _s(row.get("Deliverable")),
                        "amount_usd": float(row.get(
                            "Maximum Grant instalment payable on satisfaction of this deliverable (USD)") or 0.0),
                    })

            # --- Import Project Overview sheet (if present) and update Project Overview state ---
            if "Project Overview" in xls.sheet_names:
                id_df = pd.read_excel(xls, sheet_name="Project Overview")
                try:
                    kv = {
                        str(r["Field"]).strip(): (
                            "" if (pd.isna(r["Value"])) else str(r["Value"])
                        )
                        for _, r in id_df.iterrows()
                    }
                except Exception:
                    kv = {}


                def _g(field_label: str) -> str:
                    return (kv.get(field_label, "") or "").strip()

                id_info = st.session_state.get("id_info", {}) or {}
                id_info.update({
                    "title": _g(LABELS["title"]),
                    "pi_name": _g(LABELS["pi_name"]),
                    "pi_email": _g(LABELS["pi_email"]),
                    "implementing_partners": _g(LABELS["implementing_partners"]),
                    "supporting_partners": _g(LABELS["supporting_partners"]),
                    "start_date": parse_date_like(kv.get(LABELS["start_date"], "")) or None,
                    "end_date": parse_date_like(kv.get(LABELS["end_date"], "")) or None,
                    "location": _g(LABELS["location"]),
                    "contact_name": _g(LABELS["contact_name"]),
                    "contact_email": _g(LABELS["contact_email"]),
                    "contact_phone": _g(LABELS["contact_phone"]),
                })

                st.session_state.id_info = id_info

                # prime live widget values so inputs show imported data immediately
                st.session_state["id_title"] = id_info["title"]
                st.session_state["id_pi_name"] = id_info["pi_name"]
                st.session_state["id_pi_email"] = id_info["pi_email"]
                st.session_state["id_implementing_partners"] = id_info["implementing_partners"]
                st.session_state["id_supporting_partners"] = id_info["supporting_partners"]
                st.session_state["id_start_date"] = id_info["start_date"]
                st.session_state["id_end_date"] = id_info["end_date"]
                st.session_state["id_location"] = id_info["location"]
                st.session_state["id_contact_name"] = id_info["contact_name"]
                st.session_state["id_contact_email"] = id_info["contact_email"]
                st.session_state["id_contact_phone"] = id_info["contact_phone"]

            # Remember we loaded this file content; prevents re-import on button clicks
            st.session_state["_resume_file_sig"] = file_sig
            tabs[0].success("✅ Previous submission loaded into session.")
            st.rerun()

        # else: same file uploaded again → skip re-import so edit/delete works
    except Exception as e:
        tabs[0].error(f"Could not parse uploaded Excel: {e}")

# ===== TAB 2: Project Overview =====
with tabs[1]:
    st.header("🪪 Project Overview")

    # defaults
    if "id_info" not in st.session_state:
        st.session_state.id_info = {
            "title": "",
            "pi_name": "",
            "pi_email": "",
            "implementing_partners": "",
            "supporting_partners": "",
            "start_date": None,
            "end_date": None,
            "location": "",
            "contact_name": "",
            "contact_email": "",
            "contact_phone": "",
        }

    # ensure widget keys exist (so inputs are persistent & can be set by the importer)
    for k, v in [
        ("id_title", st.session_state.id_info["title"]),
        ("id_pi_name", st.session_state.id_info["pi_name"]),
        ("id_pi_email", st.session_state.id_info["pi_email"]),
        ("id_implementing_partners", st.session_state.id_info["implementing_partners"]),
        ("id_supporting_partners", st.session_state.id_info["supporting_partners"]),
        ("id_start_date", st.session_state.id_info["start_date"]),
        ("id_end_date", st.session_state.id_info["end_date"]),
        ("id_location", st.session_state.id_info["location"]),
        ("id_contact_name", st.session_state.id_info["contact_name"]),
        ("id_contact_email", st.session_state.id_info["contact_email"]),
        ("id_contact_phone", st.session_state.id_info["contact_phone"]),
    ]:
        if k not in st.session_state:
            st.session_state[k] = v
    st.session_state.id_info["title"] = st.text_input(LABELS["title"], key="id_title")
    st.session_state.id_info["pi_name"] = st.text_input(LABELS["pi_name"], key="id_pi_name")
    st.session_state.id_info["pi_email"] = st.text_input(LABELS["pi_email"], key="id_pi_email")
    st.session_state.id_info["implementing_partners"] = st.text_input(LABELS["implementing_partners"],
                                                                      key="id_implementing_partners")
    st.session_state.id_info["supporting_partners"] = st.text_input(LABELS["supporting_partners"],
                                                                    key="id_supporting_partners")
    st.session_state.id_info["location"] = st.text_input(LABELS["location"], key="id_location")
    st.session_state.id_info["contact_name"] = st.text_input(LABELS["contact_name"], key="id_contact_name")
    st.session_state.id_info["contact_email"] = st.text_input(LABELS["contact_email"], key="id_contact_email")
    st.session_state.id_info["contact_phone"] = st.text_input(LABELS["contact_phone"], key="id_contact_phone")

    def _as_date(val):
        return val if isinstance(val, date) else None

    if "id_start_date" in st.session_state and _as_date(st.session_state["id_start_date"]) is None:
        del st.session_state["id_start_date"]
    if "id_end_date" in st.session_state and _as_date(st.session_state["id_end_date"]) is None:
        del st.session_state["id_end_date"]

    start_default = _as_date(st.session_state.id_info.get("start_date")) or date.today()
    end_default = _as_date(st.session_state.id_info.get("end_date")) or date.today()

    date_cols = st.columns(2)
    with date_cols[0]:
        sd = st.date_input(
            "Project start date",
            value=None if "id_start_date" in st.session_state else start_default,
            key="id_start_date",
            format="DD/MM/YYYY",
        )
    with date_cols[1]:
        ed = st.date_input(
            "Project end date",
            value=None if "id_end_date" in st.session_state else end_default,
            key="id_end_date",
            format="DD/MM/YYYY",
        )

    st.session_state.id_info["start_date"] = sd if isinstance(sd, date) else None
    st.session_state.id_info["end_date"] = ed if isinstance(ed, date) else None

    # inline validation (no button)
    errs = []
    ii = st.session_state.id_info
    if not ii["title"].strip():
        errs.append("Project title is required.")
    if not ii["pi_name"].strip():
        errs.append("PI name is required.")
    if not ii["location"].strip():
        errs.append("Implementation location is required.")
    if not ii["pi_email"].strip() or "@" not in ii["pi_email"] or "." not in ii["pi_email"]:
        errs.append("Valid PI email is required.")
    if ii["contact_email"].strip() and ("@" not in ii["contact_email"] or "." not in ii["contact_email"]):
        errs.append("Main Contact email must be valid if provided.")
    if ii["start_date"] and ii["end_date"] and ii["start_date"] > ii["end_date"]:
        errs.append("Project start date must be on or before the end date.")

    # show errors only after the user has typed something in any required field
    touched = any([
        st.session_state["id_title"].strip(),
        st.session_state["id_pi_name"].strip(),
        st.session_state["id_pi_email"].strip(),
        st.session_state["id_supporting_partners"].strip(),
        st.session_state["id_location"].strip(),
    ])
    if touched:
        for e in errs:
            st.error(e)


    # ---------- Read-only computed fields ----------

    budget_total = sum(float(r.get("total_usd") or 0.0) for r in st.session_state.get("budget", []))
    outputs_count = len(st.session_state.get("outputs", []))
    kpis_count = len(st.session_state.get("kpis", []))
    activities_count = len(st.session_state.get("workplan", []))

    # ---------- Summary cards (dashboard-style) ----------

    def format_k_usd(val: float) -> str:
        if val >= 1000:
            return f"USD {val / 1000:,.0f}k"
        return f"USD {val:,.0f}"

    # (optional) project duration in months as an extra metric
    start = st.session_state.get("id_info", {}).get("start_date")
    end = st.session_state.get("id_info", {}).get("end_date")
    duration_months = 0
    if start and end:
        duration_months = max(0, int((end - start).days / 30.4375))  # avg month length

    # Card CSS (scoped to avoid collisions)
    st.markdown("""
    <style>
    .cards-grid {
      display: grid;
      grid-template-columns: repeat(2, minmax(280px, 1fr));
      gap: 18px;
      margin: 10px 0 8px;
    }
    @media (min-width: 1024px) {
      .cards-grid { grid-template-columns: repeat(4, minmax(250px, 1fr)); }
    }
    .card {
      background: #ffffff;
      border-radius: 16px;
      box-shadow: 0 6px 18px rgba(16,24,40,.06);
      padding: 18px 18px 20px;
      border: 1px solid rgba(16,24,40,.06);
      position: relative;
    }
    .card h4 {
      margin: 0;
      font-size: 1.05rem;
      font-weight: 600;
      color: #1f2937;
    }
    .card .value {
      font-size: 2rem;
      font-weight: 800;
      margin-top: 6px;
      color: #0f172a;
      letter-spacing: -0.02em;
    }
    .card .badge {
      position: absolute;
      right: 14px;
      top: 14px;
      width: 42px; height: 42px;
      border-radius: 12px;
      display: grid; place-items: center;
      color: #0f172a;
      font-size: 20px;
    }
    .badge-blue   { background: rgba(59,130,246,.12); color:#2563eb; }
    .badge-green  { background: rgba(16,185,129,.12); color:#059669; }
    .badge-indigo { background: rgba(99,102,241,.12); color:#4f46e5; }
    .badge-slate  { background: rgba(148,163,184,.18); color:#334155; }
    </style>
    """, unsafe_allow_html=True)

    st.markdown("### Summary", unsafe_allow_html=True)

    st.markdown(f"""
    <div class="cards-grid">

      <div class="card">
        <div class="badge badge-blue">💵</div>
        <h4>Total Funding requested</h4>
        <div class="value">{format_k_usd(budget_total)}</div>
      </div>

      <div class="card">
        <div class="badge badge-indigo">📦</div>
        <h4># Outputs</h4>
        <div class="value">{outputs_count:,}</div>
      </div>

      <div class="card">
        <div class="badge badge-green">🗂️</div>
        <h4># Activities</h4>
        <div class="value">{activities_count:,}</div>
      </div>

      <div class="card">
        <div class="badge badge-slate">🎯</div>
        <h4># KPIs</h4>
        <div class="value">{kpis_count:,}</div>
      </div>

    </div>
    """, unsafe_allow_html=True)

    # Optional: a small note below the cards
    st.caption(footer_note)

    # Cross-check note area (non-blocking warnings placeholder)
    warnings = []
    # (Optional) You can later add KPI due-date vs project dates checks here and append messages to `warnings`.
    if warnings:
        for w in warnings:
            st.warning(w)


# ===== TAB 3: Logframe =====
tabs[2].header("📊 Build Your Logframe")
inject_logframe_css()
# --- numbering for labels shown in UI and preview ---
out_nums, kpi_nums = compute_numbers()   # outputs -> 'n', KPIs -> 'n.p'

# --- Add forms ---
with tabs[2].expander("➕ Add Goal"):
    if len(st.session_state.impacts) >= 1:
        st.info("Only one Goal is allowed. Edit the existing Goal in the preview below.")
    else:
        with st.form("goal_form"):
            goal_text = st.text_area("Goal (single, high-level statement)")
            goal_assumptions = st.text_area("Key Assumptions (optional)")
            if st.form_submit_button("Add Goal") and goal_text.strip():
                st.session_state.impacts.append({
                    "id": generate_id(),
                    "level": "Goal",
                    "name": goal_text.strip(),
                    "assumptions": goal_assumptions.strip(),
                })

with tabs[2].expander("➕ Add Outcome"):
    if not st.session_state.impacts:
        st.warning("Add the Goal first.")
    elif len(st.session_state.outcomes) >= 1:
        st.info("Only one Outcome is allowed. Edit the existing Outcome in the preview below.")
    else:
        with st.form("outcome_form"):
            outcome_text = st.text_area("Outcome (statement)")
            # since there is only one goal, no need to pick it; link to the single goal
            linked_goal_id = st.session_state.impacts[0]["id"]
            if st.form_submit_button("Add Outcome") and outcome_text.strip():
                st.session_state.outcomes.append(
                    {"id": generate_id(), "level": "Outcome", "name": outcome_text.strip(), "parent_id": linked_goal_id}
                )

# --- Add ONE Outcome-level KPI ---
with tabs[2].expander("➕ Add Outcome KPI (one per outcome)"):
    if not st.session_state.outcomes:
        tabs[2].warning("Add the Outcome first.")
    else:
        outcome = st.session_state.outcomes[0]   # single-outcome design
        existing_outcome_kpis = [
            k for k in st.session_state.kpis
            if k.get("parent_level") == "Outcome" and k.get("parent_id") == outcome["id"]
        ]

        if existing_outcome_kpis:
            st.info("This outcome already has a KPI. Edit it below in the preview.")
        else:
            with st.form("kpi_outcome_form"):
                kpi_text = st.text_area("Outcome KPI*")
                baseline = st.text_input("Baseline")
                target   = st.text_input("Target")
                payment_linked = st.checkbox("Linked to Payment (optional)")
                mov = st.text_area("Means of Verification")

                if st.form_submit_button("Add Outcome KPI") and kpi_text.strip():
                    st.session_state.kpis.append({
                        "id": generate_id(),
                        "level": "KPI",
                        "name": strip_label_prefix(kpi_text.strip(), "KPI"),
                        "parent_id": outcome["id"],
                        "parent_level": "Outcome",   # <<< key difference
                        "baseline": baseline.strip(),
                        "target": target.strip(),
                        "linked_payment": bool(payment_linked),
                        "mov": mov.strip(),
                    })
                    st.rerun()

with tabs[2].expander("➕ Add Output"):
    if not st.session_state.outcomes:
        tabs[2].warning("Add the Outcome first.")
    else:
        with st.form("output_form"):
            output_title = st.text_input("Output title (e.g., 'Output 1')")
            output_assumptions = st.text_area("Key Assumptions (optional)")
            if st.form_submit_button("Add Output") and output_title.strip():
                linked_outcome_id = st.session_state.outcomes[0]["id"]
                st.session_state.outputs.append(
                    {
                        "id": generate_id(),
                        "level": "Output",
                        "name": output_title.strip(),
                        "parent_id": linked_outcome_id,
                        "assumptions": output_assumptions.strip(),
                    }
                )

with tabs[2].expander("➕ Add KPI"):
    if not st.session_state.outputs:
        tabs[2].warning("Add an Output first.")
    else:
        with st.form("kpi_form"):
            parent = st.selectbox(
                "Parent Output",
                st.session_state.outputs,
                format_func=lambda o: f"Output {out_nums.get(o['id'],'?')} — {o.get('name','Output')}"
            )
            kpi_text = st.text_area("KPI*")
            baseline = st.text_input("Baseline")
            target   = st.text_input("Target")
            payment_linked = st.checkbox("Linked to Payment (optional)")
            mov = st.text_area("Means of Verification")

            if st.form_submit_button("Add KPI") and kpi_text.strip():
                st.session_state.kpis.append({
                    "id": generate_id(),
                    "level": "KPI",
                    "name": strip_label_prefix(kpi_text.strip(), "KPI"),
                    "parent_id": parent["id"],
                    "parent_level": "Output",   # <-- fixed
                    "baseline": baseline.strip(),
                    "target": target.strip(),
                    "linked_payment": bool(payment_linked),
                    "mov": mov.strip(),
                })
                st.rerun()

# ---- View helpers (card layout, compact-aware) ----
def view_goal(g):
    title = escape(g.get("name", ""))
    header_html = f"<div class='lf-goal-header'><strong>Goal:</strong> {title}</div>"
    ass_html = _assumptions_html(g.get("assumptions"))
    return view_logframe_element(header_html + ass_html, kind="goal")

def view_outcome(o):
    return f"##### 🟪 **Outcome:** {o.get('name','')}"

def view_output(out):
    num   = out_nums.get(out["id"], "?")
    title = out.get("name", "Output")

    # header line (keeps your green card style via view_logframe_element)
    header_html = (
        f"<div class='lf-out-header'><strong>Output {num}:</strong> {title}</div>"
    )

    # assumptions -> bullet list (one bullet per line the user typed)
    ass_html = _assumptions_html(out.get("assumptions"))

    # wrap in the green card
    return view_logframe_element(header_html + ass_html, kind="output")

def view_output_header(out):
    num   = out_nums.get(out["id"], "?")
    title = escape(out.get("name", "Output"))
    header_html = f"<div class='lf-out-header'><strong>Output {num}:</strong> {title}</div>"
    return view_logframe_element(header_html, kind="output")

def view_kpi(k):
    # decide label prefix based on level
    is_outcome = (k.get("parent_level") == "Outcome")
    num  = kpi_nums.get(k["id"], "?") if not is_outcome else ""  # no numbering for outcome-level
    name = k.get("name", "")

    bp  = (k.get("baseline") or "").strip()
    tg  = (k.get("target") or "").strip()
    mov = (k.get("mov") or "").strip()

    chip = (
        "<span class='chip green'>Payment-linked</span>"
        if k.get("linked_payment")
        else "<span class='chip'>Not payment-linked</span>"
    )

    # Title
    if is_outcome:
        header = f"<div class='lf-kpi-title'>Outcome KPI: {name}</div>"
    else:
        header = f"<div class='lf-kpi-title'>KPI {num}: {name}</div>"

    lines = []
    if bp:
        lines.append(f"<div class='lf-line'><b>Baseline:</b> {bp}</div>")
    if tg:
        lines.append(f"<div class='lf-line'><b>Target:</b> {tg}</div>")
    if mov:
        lines.append(f"<div class='lf-line'><b>Means of Verification:</b> {mov}</div>")

    lines.append(f"<div class='lf-line'>{chip}</div>")
    inner = header + "".join(lines)
    return view_logframe_element(inner, kind="kpi")

def view_activity(a: dict, act_label: str, id_to_output: dict, id_to_kpi: dict) -> str:
    """
    Render one activity as a card.
    - a: activity dict from st.session_state.workplan
    - act_label: precomputed label like "1.2"
    - id_to_output: {output_id -> output name}
    - id_to_kpi:    {kpi_id -> kpi name}
    """
    out_name = id_to_output.get(a.get("output_id"), "(unassigned)")
    title    = f"Activity {escape(act_label)} — {escape(a.get('name',''))}"
    owner    = escape(a.get("owner","") or "—")
    sd       = fmt_dd_mmm_yyyy(a.get("start")) or "—"
    ed       = fmt_dd_mmm_yyyy(a.get("end"))   or "—"
    kpis_txt = ", ".join(escape(id_to_kpi.get(kid,"")) for kid in (a.get("kpi_ids") or [])) or "—"

    title_html = f"<div class='lf-activity-title'>{title}</div>"
    rows = [
        ("Output", out_name),
        ("Owner", owner),
        ("Start date", sd),
        ("End date", ed),
        ("Linked KPIs", kpis_txt),
    ]

    # Build rows (simple label/value pairs). You can skip empty ones here if desired.
    body = "".join(
        f"<div class='lf-line'><b>{escape(lbl)}:</b> {val}</div>"
        for (lbl, val) in rows
        if (val is not None and str(val).strip() != "")
    )

    # Use your generic card wrapper (orange indicators style)
    return view_logframe_element(title_html + body, kind="activity")

def make_ext_id(kind: str, text: str) -> str:
    """
    Deterministic short id for external linking across exports/imports.
    kind: 'output' | 'kpi'
    text: any stable text (e.g., output name; for KPI: 'OutputName|KPI text')
    """
    base = f"{kind}:{(text or '').strip().lower()}"
    return hashlib.sha1(base.encode("utf-8")).hexdigest()[:10]

# --- Inline preview with Edit / Delete buttons (refactored, card layout) ---
with tabs[2]:
    st.markdown("---")
    st.subheader("Current Logframe (preview) — click ✏️ to edit, 🗑️ to delete")

    for g in st.session_state.get("impacts", []):
        render_editable_item(
            item=g,
            list_name="impacts",
            edit_flag_key="edit_goal",
            view_md_func=view_goal,
            fields=[
                ("name", st.text_area, "Goal"),
                ("assumptions", st.text_area, "Key Assumptions"),
            ],
            on_delete=lambda _id=g["id"]: (delete_cascade(goal_id=_id), st.rerun()),
            key_prefix="lf"
        )

        outcomes_here = [o for o in st.session_state.get("outcomes", []) if o.get("parent_id") == g["id"]]
        for oc in outcomes_here:
            render_editable_item(
                item=oc,
                list_name="outcomes",
                edit_flag_key="edit_outcome",
                view_md_func=view_outcome,
                default_label="Outcome",
                on_delete=lambda _id=oc["id"]: (delete_cascade(outcome_id=_id), st.rerun()),
                key_prefix="lf"
            )

            # --- Outcome-level KPI preview (edit/delete same as output KPI) ---
            outcome_kpis = [k for k in st.session_state.get("kpis", [])
                            if k.get("parent_id") == oc["id"] and k.get("parent_level") == "Outcome"]

            for k in outcome_kpis:
                render_editable_item(
                    item=k, list_name="kpis", edit_flag_key="edit_kpi",
                    view_md_func=lambda kk: view_kpi(kk),  # reuse the same renderer
                    fields=[
                        ("name", st.text_area, "KPI"),
                        ("baseline", st.text_input, "Baseline"),
                        ("target", st.text_input, "Target"),
                        ("linked_payment",
                         lambda label, value, key: st.checkbox(label, value=bool(value), key=key),
                         "Linked to Payment"),
                        ("mov", st.text_area, "Means of Verification"),
                    ],
                    on_delete=lambda _id=k["id"]: (
                        setattr(st.session_state, "kpis", [x for x in st.session_state.kpis if x["id"] != _id]),
                        st.rerun()
                    ),
                    key_prefix="lf"
                )

            outs_here = [o for o in st.session_state.get("outputs", []) if o.get("parent_id") == oc["id"]]
            for out in outs_here:
                with st.container():  # now this container lives inside the Logframe tab
                    render_editable_item(
                        item=out,
                        list_name="outputs",
                        edit_flag_key="edit_output",
                        view_md_func=view_output,
                        fields=[
                            ("name", st.text_input, "Output title"),
                            ("assumptions", st.text_area, "Key Assumptions"),
                        ],
                        on_delete=lambda _id=out["id"]: (delete_cascade(output_id=_id), st.rerun()),
                        key_prefix="lf"
                    )

                    k_children = [k for k in st.session_state.get("kpis", [])
                                  if (k.get("parent_id") == out["id"])]  # parent_level check optional if you dropped outcome-level KPIs
                    for k in k_children:
                        render_editable_item(
                            item=k, list_name="kpis", edit_flag_key="edit_kpi",
                            view_md_func=view_kpi,
                            fields=[
                                ("parent_id",
                                 lambda label, value, key: select_output_id(label, value, key),
                                 "Linked Output"),                                ("name", st.text_area, "KPI"),
                                ("baseline", st.text_input, "Baseline"),
                                ("target", st.text_input, "Target"),
                                ("linked_payment",
                                 lambda label, value, key: st.checkbox(label, value=bool(value), key=key),
                                 "Linked to Payment"),
                                ("mov", st.text_area, "Means of Verification"),
                            ],
                            on_delete=lambda _id=k["id"]: (
                                setattr(st.session_state, "kpis", [x for x in st.session_state.kpis if x["id"] != _id]),
                                st.rerun()
                            ),
                            key_prefix="lf"
                        )

# ===== TAB 4: Workplan =====
with tabs[3]:
    st.header("📆 Workplan")
    with st.expander("➕ Add Activity"):
        with st.form("workplan_form_v2"):
            # Required: link to Output
            output_parent = st.selectbox(
                "Linked Output*",
                st.session_state.outputs,
                format_func=lambda x: x.get("name") or "Output"
            )

            # Optional: link to KPI(s) under that Output
            output_id = output_parent["id"] if output_parent else None
            kpis_for_output = [k for k in st.session_state.kpis if k.get("parent_id") == output_id]
            kpi_links = st.multiselect(
                "Linked KPI(s) (optional)",
                kpis_for_output,
                format_func=lambda k: f"{k.get('name','')}"
            )

            name = st.text_input("Activity*")
            owner = st.text_input("Responsible person/institution*")
            c1, c2 = st.columns(2)
            with c1:
                start = st.date_input("Start date*")
            with c2:
                end = st.date_input("End date*")

            submitted = st.form_submit_button("Add to Workplan")
            if submitted and name.strip() and owner.strip() and output_parent and start and end and start <= end:
                st.session_state.workplan.append({
                    "id": generate_id(),
                    "output_id": output_id,
                    "name": name.strip(),
                    "kpi_ids": [k["id"] for k in kpi_links],
                    "owner": owner.strip(),
                    "start": start,
                    "end": end
                })
                st.rerun()
            elif submitted:
                st.warning("Please fill required fields (Output, Activity, Owner, Start≤End).")

    # ------- Gantt (web) -------
    df_g = _workplan_df()
    with st.expander("📈 Gantt chart (Workplan)", expanded=bool(len(df_g))):
        if df_g.empty:
            st.info("Add activities with start & end dates to see the Gantt.")
        else:
            height = min(1.2 + 0.35 * len(df_g), 12)
            fig, ax = plt.subplots(figsize=(11, height))
            _draw_gantt(ax, df_g, show_today=False)
            fig.subplots_adjust(left=0.26, right=0.98, top=0.96, bottom=0.30)
            st.pyplot(fig, clear_figure=True, use_container_width=True)

    # --- Card view (optional: put this below or instead of the table) ---
    out_nums, kpi_nums, act_nums = compute_numbers(include_activities=True)
    id_to_output = {o["id"]: (o.get("name") or "Output") for o in st.session_state.outputs}
    id_to_kpi    = {k["id"]: (k.get("name") or "")       for k in st.session_state.kpis}

    for oc in st.session_state.outcomes:
        outs_here = [o for o in st.session_state.outputs if o.get("parent_id") == oc["id"]]
        for out in outs_here:
            # green Output header card
            st.markdown(view_output_header(out), unsafe_allow_html=True)

            # orange Activity cards (with edit/delete)
            acts_here = [a for a in st.session_state.workplan if a.get("output_id") == out["id"]]
            for a in acts_here:
                label = act_nums.get(a["id"], "?")
                # Edit mode?
                if st.session_state.get("edit_activity") == a["id"]:
                    e1, e2, e3 = st.columns([0.90, 0.05, 0.05])
                    with e1:
                        new_output_id = select_output_id("Linked Output", a.get("output_id"), f"a_out_{a['id']}")
                        new_name = st.text_input("Activity", value=a.get("name", ""), key=f"a_name_{a['id']}")
                        new_owner = st.text_input("Owner", value=a.get("owner", ""), key=f"a_owner_{a['id']}")
                        cA, cB = st.columns(2)
                        with cA:
                            new_start = st.date_input("Start date", value=a.get("start"), key=f"a_start_{a['id']}")
                        with cB:
                            new_end = st.date_input("End date", value=a.get("end"), key=f"a_end_{a['id']}")
                    if e2.button("💾", key=f"a_save_{a['id']}"):
                        idx = _find_by_id(st.session_state.workplan, a["id"])
                        if idx is not None:
                            st.session_state.workplan[idx].update({
                                "output_id": new_output_id,
                                "name": new_name.strip(),
                                "owner": new_owner.strip(),
                                "start": new_start,
                                "end": new_end,
                            })
                        st.session_state["edit_activity"] = None
                        st.rerun()
                    if e3.button("✖️", key=f"a_cancel_{a['id']}"):
                        st.session_state["edit_activity"] = None
                        st.rerun()
                else:
                    v1, v2, v3 = st.columns([0.90, 0.05, 0.05])
                    v1.markdown(
                        view_activity_readonly(a, label, id_to_output, id_to_kpi),
                        unsafe_allow_html=True
                    )
                    if v2.button("✏️", key=f"a_edit_{a['id']}"):
                        st.session_state["edit_activity"] = a["id"]
                        st.rerun()
                    if v3.button("🗑️", key=f"a_del_{a['id']}"):
                        st.session_state.workplan = [x for x in st.session_state.workplan if x["id"] != a["id"]]
                        st.rerun()

# ===== TAB 5: Budget =====
with tabs[4]:
    st.header("💵 Define Budget")
    st.caption("Enter amounts in USD")

    # ---------- session state ----------
    if "edit_budget_row" not in st.session_state:
        st.session_state.edit_budget_row = None

    # ---------- lookups ----------
    activities = st.session_state.get("workplan", [])
    out_nums, _, act_nums = compute_numbers(include_activities=True)
    act_lookup = activity_label_map(act_nums)

    def _activity_label(aid: str | None) -> str:
        return act_lookup.get(aid, "(none)")

    # keep Budget “containers” in sync with the Workplan
    from collections import defaultdict
    st.session_state.setdefault("_budget_activity_ids", set())
    st.session_state["_budget_activity_ids"] = {a["id"] for a in st.session_state.get("workplan", [])}

    # group existing budget lines by activity id
    lines_by_act = defaultdict(list)
    for row in st.session_state.get("budget", []):
        lines_by_act[row.get("activity_id")].append(row)

    # ---------- CSS (small tweaks) ----------
    st.markdown("""
    <style>
      div[data-testid="column"] label p { font-weight:600; }
      .stSelectbox div[role="combobox"] span,
      .stSelectbox [role="listbox"] div { white-space:normal !important; line-height:1.3 !important; }
      .stSelectbox div[aria-haspopup="listbox"]>div { white-space:normal !important; }
      .budget-cell{white-space:normal !important; line-height:1.35;}
      .total-chip{
        font-weight:700; padding:.30rem .55rem; background:#fff;
        border-radius:6px; border:1px solid rgba(0,0,0,.10);
        display:inline-block; white-space:nowrap;
      }
    </style>
    """, unsafe_allow_html=True)

    # =========================================================
    # GROUPED VIEW (Activity → line items → subtotal) + per-activity add form
    # =========================================================

    grand_total = 0.0

    for a in st.session_state.get("workplan", []):
        if a["id"] not in st.session_state["_budget_activity_ids"]:
            continue

        # Activity header (inside yellow box)
        st.markdown(
            view_logframe_element(
                f"<div class='lf-activity-title'>{escape(format_activity_label(a, act_nums))}</div>",
                kind="activity"  # uses lf-card lf-card--activity
            ),
            unsafe_allow_html=True,
        )

        # Budget table header (placed just below yellow box)
        h1, h2, h3, h4, h5, h6, h7, h8 = st.columns([0.28, 0.18, 0.18, 0.12, 0.10, 0.06, 0.06, 0.08])
        h1.markdown("**Line Item**")
        h2.markdown("**Category**")
        h3.markdown("**Sub Category**")
        h4.markdown("**Unit**")
        h5.markdown("**Unit Cost (USD)**")
        h6.markdown("**Qty**")
        h7.markdown("**Total (USD)**")
        h8.markdown("**Actions**")

        # Existing lines under this activity
        subtotal = 0.0
        for r in lines_by_act.get(a["id"], []):
            rid = r["id"]

            # inline editor FIRST
            if st.session_state.get("edit_budget_row") == rid:
                _render_budget_row_editor(r, rid, act_lookup)
                continue

            # --- Static row (one line item) ---
            l1, l2, l3, l4, l5, l6, l7, l8 = st.columns([0.28, 0.18, 0.18, 0.12, 0.10, 0.06, 0.06, 0.08])

            # values
            item = r.get("item", "—")
            cat = r.get("category", "—")
            subc = r.get("subcategory", "—")
            unit = r.get("unit", "—")
            uc = float(r.get("unit_cost") or 0.0)
            qty = float(r.get("qty") or 0.0)
            tot = float(r.get("total_usd") or (uc * qty) or 0.0)

            # cells
            l1.markdown(f"<div class='budget-cell'>{escape(item)}</div>", unsafe_allow_html=True)
            l2.write(cat)
            l3.write(subc)
            l4.write(unit)
            l5.write(f"{uc:,.2f}")
            l6.write(f"{qty:,.2f}")
            l7.write(f"{tot:,.2f}")

            # actions
            b_edit, b_del = l8.columns(2)
            if b_edit.button("✏️", key=f"edit_{rid}", help="Edit line"):
                st.session_state.edit_budget_row = rid
                st.session_state[f"e_init_{rid}"] = False  # seed-once flag for the editor
                st.rerun()
            if b_del.button("🗑️", key=f"del_{rid}", help="Delete line"):
                st.session_state.budget = [x for x in st.session_state.budget if x["id"] != rid]
                st.rerun()

            # accumulate subtotal
            subtotal += tot

        st.markdown(
            f"<div style='text-align:right; font-weight:700;'>Subtotal: USD {subtotal:,.2f}</div>",
            unsafe_allow_html=True,
        )

        grand_total += subtotal

            # -------- Add new line to THIS activity (uses dynamic unit mapping) --------
        with st.expander("➕ Add Budget Line to this Activity"):
            form_prefix = f"add_{a['id']}"
            item_key       = f"{form_prefix}_item"
            cat_key        = f"{form_prefix}_cat"
            sub_key        = f"{form_prefix}_sub"
            unit_choice_key   = f"{form_prefix}_unit_choice"
            unit_custom_key   = f"{form_prefix}_unit_custom"
            unit_prev_sub_key = f"{form_prefix}_unit_prev_sub"
            unit_reset_key    = f"{form_prefix}_unit_reset"

            # Category
            cat_options = list(CATEGORY_TREE.keys())
            st.session_state.setdefault(cat_key, cat_options[0])
            current_cat = st.session_state[cat_key]

            # Sub-category (dependent)
            sub_options = subcategories_for(current_cat) or ["(none)"]
            if sub_key not in st.session_state or st.session_state[sub_key] not in sub_options:
                st.session_state[sub_key] = sub_options[0] if sub_options else ""
            current_sub = st.session_state[sub_key]

            # Unit options (from your mapping)
            unit_values, unit_required = unit_choices_for(current_cat, current_sub)
            allow_custom = not unit_required

            # Build unit choices list
            unit_options = ["Select unit"] + list(dict.fromkeys(unit_values))  # de-dupe, preserve order
            if allow_custom:
                unit_options.append("Custom…")

            st.session_state.setdefault(unit_choice_key, "Select unit")
            st.session_state.setdefault(unit_custom_key, "")
            st.session_state.setdefault(unit_prev_sub_key, current_sub)
            st.session_state.setdefault(unit_reset_key, False)

            # Reset unit if subcategory changed and previous selection is invalid
            if st.session_state[unit_prev_sub_key] != current_sub:
                sel = st.session_state[unit_choice_key]
                if sel not in unit_options or (sel == "Custom…" and not allow_custom):
                    st.session_state[unit_choice_key] = "Select unit"
                    st.session_state[unit_custom_key] = ""
                    st.session_state[unit_reset_key]  = unit_required
                else:
                    st.session_state[unit_reset_key] = False
                st.session_state[unit_prev_sub_key] = current_sub

            # Form rows
            item = st.text_input("Line Item*", key=item_key)

            r2_c1, r2_c2, r2_c3 = st.columns(3)
            current_cat = r2_c1.selectbox("Cost Category*", cat_options,
                                          index=cat_options.index(st.session_state[cat_key]), key=cat_key)

            sub_options = subcategories_for(current_cat) or ["(none)"]
            if st.session_state[sub_key] not in sub_options:
                st.session_state[sub_key] = sub_options[0] if sub_options else ""
            current_sub = r2_c2.selectbox("Sub Category*", sub_options,
                                          index=sub_options.index(st.session_state[sub_key]), key=sub_key)

            # refresh unit choices after any category/sub change
            unit_values, unit_required = unit_choices_for(current_cat, current_sub)
            allow_custom = not unit_required
            unit_options = ["Select unit"] + list(dict.fromkeys(unit_values))
            if allow_custom:
                unit_options.append("Custom…")
            unit_label = "Unit*" if unit_required else "Unit"
            chosen = r2_c3.selectbox(unit_label, unit_options,
                                     index=unit_options.index(st.session_state[unit_choice_key]),
                                     key=unit_choice_key)
            custom_unit = ""
            if chosen == "Custom…" and allow_custom:
                custom_unit = r2_c3.text_input(" ", key=unit_custom_key)
            elif chosen != "Custom…":
                st.session_state[unit_custom_key] = ""

            if unit_required and st.session_state.get(unit_reset_key):
                r2_c3.caption("Unit reset — please choose a unit for this sub-category.")

            # Cost & quantity
            cL, cR = st.columns(2)
            uc  = cL.number_input("Unit Cost (USD)*", min_value=0.0, value=0.0, step=10.0, format="%.2f",
                                  key=f"{form_prefix}_uc")
            qty = cR.number_input("Qty*", min_value=0.0, value=1.0, step=1.0, format="%.2f",
                                  key=f"{form_prefix}_qty")
            line_total = float(uc) * float(qty)

            # Determine final unit string
            if chosen == "Select unit":
                final_unit = ""
            elif chosen == "Custom…" and allow_custom:
                final_unit = (custom_unit or "").strip()
            else:
                final_unit = chosen

            left, right = st.columns([0.20, 0.80])
            if left.button("Add line", key=f"add_line_btn_{a['id']}"):
                if not item.strip():
                    st.warning("Line Item is required.")
                elif not current_cat or not current_sub or current_sub == "(none)":
                    st.warning("Pick a Cost Category and Sub Category.")
                elif unit_required and not final_unit:
                    st.warning("Please choose a Unit for this Sub Category.")
                elif uc <= 0 or qty <= 0:
                    st.warning("Unit cost and quantity must be positive.")
                else:
                    st.session_state.budget.append({
                        "id": generate_id(),
                        "activity_id": a["id"],
                        "item": item.strip(),
                        "category": current_cat,
                        "subcategory": current_sub,
                        "unit": final_unit,
                        "unit_cost": float(uc),
                        "qty": float(qty),
                        "currency": "USD",
                        "total_usd": line_total,
                    })
                    st.session_state[unit_reset_key] = False
                    st.rerun()

            right.markdown(
                f"<div style='text-align:right;'><span class='total-chip'>USD {line_total:,.2f}</span></div>",
                unsafe_allow_html=True,
            )

    # Grand total for the whole budget
    st.markdown(f"**Total: USD {grand_total:,.2f}**")
    st.caption(footer_note)

# ===== TAB 6: Disbursement Schedule =====
with tabs[5]:
    st.header("💸 Disbursement Schedule")

    # Keep this table derived from KPIs (add/update/remove) while preserving date/amount
    sync_disbursement_from_kpis()

    if not st.session_state.disbursement:
        st.info("No KPIs are marked as **Linked to Payment** in the Logframe.")
    else:
        out_nums, _ = compute_numbers()
        id_to_output = {o["id"]: (o.get("name") or "Output") for o in st.session_state.outputs}

        def _out_label(d):
            n = out_nums.get(d.get("output_id",""), "")
            return f"{n} | {id_to_output.get(d.get('output_id'), '(unassigned)')}"

        st.caption("This table lists only KPIs marked as *Linked to Payment* in the Logframe. Enter the anticipated date and amount (USD); KPI details are fixed.")
        # Render (sorted for stability)
        for row in sorted(st.session_state.disbursement, key=lambda x: (_out_label(x), x.get("kpi_name",""))):
            kpid = row["kpi_id"]
            with st.container():
                # Output label (disabled), Linked-KPI (disabled), Date (editable), Amount (editable)
                c1, c2, c3, c4 = st.columns([0.25, 0.35, 0.20, 0.20])
                c1.text_input("Output", value=_out_label(row), key=f"dsp_out_{kpid}", disabled=True)
                # make Linked-KPI explicitly read-only and visually greyed out
                c2.text_input("Linked KPI", value=row.get("kpi_name",""), key=f"dsp_kpi_{kpid}", disabled=True)
                new_date = c3.date_input("Anticipated date", value=row.get("anticipated_date"), key=f"dsp_date_{kpid}")
                new_amt  = c4.number_input("Amount (USD)", min_value=0.0, value=float(row.get("amount_usd") or 0.0),
                                           step=1000.0, key=f"dsp_amt_{kpid}")

                # Persist only editable fields (date, amount)
                row["anticipated_date"] = new_date
                row["amount_usd"] = float(new_amt)
    st.caption(footer_note)

# ===== TAB 7: Export =====
tabs[6].header("📤 Export Your Application")
if tabs[6].button("Generate Backup File (Excel)"):
    wb = Workbook()
    if "Sheet" in wb.sheetnames:
        wb.remove(wb["Sheet"])

    # --- Sheet 0: Project Overview ---
    def _sum_budget_for_export():
        return sum(float(r.get("total_usd") or 0.0) for r in st.session_state.get("budget", []))

    id_info = st.session_state.get("id_info", {}) or {}

    proj_title = id_info.get("title", "")
    pi_name = id_info.get("pi_name", "")
    pi_email = id_info.get("pi_email", "")
    implementing_partners = id_info.get("implementing_partners", "")
    supporting_partners = id_info.get("supporting_partners", "")
    start_date = id_info.get("start_date", "")
    end_date = id_info.get("end_date", "")
    location = id_info.get("location", "")
    contact_name = id_info.get("contact_name", "")
    contact_mail = id_info.get("contact_email", "")
    contact_phone = id_info.get("contact_phone", "")

    # live computed
    budget_total = _sum_budget_for_export()
    outputs_count = len(st.session_state.get("outputs", []))
    kpis_count = len(st.session_state.get("kpis", []))
    activities_count = len(st.session_state.get("workplan", []))

    ws_id = wb.create_sheet("Project Overview", 0)  # put it first
    ws_id.append(["Field", "Value"])
    ws_id.append([LABELS["title"], proj_title])
    ws_id.append([LABELS["pi_name"], pi_name])
    ws_id.append([LABELS["pi_email"], pi_email])
    ws_id.append([LABELS["implementing_partners"], implementing_partners])
    ws_id.append([LABELS["supporting_partners"], supporting_partners])
    ws_id.append([LABELS["start_date"], fmt_dd_mmm_yyyy(start_date)])
    ws_id.append([LABELS["end_date"], fmt_dd_mmm_yyyy(end_date)])
    ws_id.append([LABELS["location"], location])
    ws_id.append([LABELS["contact_name"], contact_name])
    ws_id.append([LABELS["contact_email"], contact_mail])
    ws_id.append([LABELS["contact_phone"], contact_phone])

    # read-only summary values
    ws_id.append([LABELS["total_funding"], f"USD {budget_total:,.2f}"])
    ws_id.append([LABELS["outputs_count"], outputs_count])
    ws_id.append([LABELS["kpis_count"], kpis_count])
    ws_id.append([LABELS["activities_count"], activities_count])

    # Sheet 1: Summary (Goal/Outcome/Output) — with explicit IDs
    s1 = wb.create_sheet("Summary", 1)
    s1.append(["RowLevel", "ID", "ParentID", "Text / Title", "Assumptions"])

    # Goal rows
    for row in st.session_state.get("impacts", []):
        s1.append(["Goal", row.get("id", ""), "", row.get("name", ""), row.get("assumptions", "")])

    # Outcome rows
    for row in st.session_state.get("outcomes", []):
        s1.append([
            "Outcome",
            row.get("id", ""),
            row.get("parent_id", ""),
            row.get("name", ""),
            row.get("assumptions", ""),
        ])

    # Output rows (include assumptions)
    for row in st.session_state.get("outputs", []):
        s1.append([
            "Output",
            row.get("id", ""),
            row.get("parent_id", ""),
            row.get("name", ""),
            row.get("assumptions", "")
        ])

    # Sheet 2: KPI Matrix — include KPIID and ParentID (labels are just helpers)
    s2 = wb.create_sheet("KPI Matrix")
    s2.append([
        "KPIID", "Parent Level", "ParentID", "Parent (label)",
        "KPI", "Baseline", "Target",
        "Linked to Payment", "Means of Verification"
    ])

    out_nums, kpi_nums = compute_numbers()
    output_title = {o["id"]: (o.get("name") or "Output") for o in st.session_state.outputs}
    outcome_name = (st.session_state.outcomes[0]["name"] if st.session_state.outcomes else "")

    for k in st.session_state.kpis:
        plevel = k.get("parent_level", "Output")
        pid = k.get("parent_id", "")
        parent_label = (
            f"Outcome — {outcome_name}" if plevel == "Outcome"
            else f"Output {out_nums.get(pid, '')} — {output_title.get(pid, '')}"
        )
        s2.append([
            k.get("id", ""),
            plevel,
            pid,
            parent_label,
            k.get("name", ""),
            k.get("baseline", ""),
            k.get("target", ""),
            "Yes" if k.get("linked_payment") else "No",
            k.get("mov", ""),
        ])

    # Workplan (export)
    out_nums, kpi_nums, act_nums = compute_numbers(include_activities=True)
    ws2 = wb.create_sheet("Workplan")
    ws2.append([
        "Activity ID", "Activity #",
        "OutputID", "Output",
        "Activity", "Owner", "Start", "End",
        "Linked KPI IDs", "Linked KPIs"
    ])

    id_to_output = {o["id"]: (o.get("name") or "Output") for o in st.session_state.outputs}
    id_to_kpi = {k["id"]: (k.get("name") or "") for k in st.session_state.kpis}

    for a in st.session_state.workplan:
        ws2.append([
            a.get("id", ""),
            act_nums.get(a["id"], ""),
            a.get("output_id", ""),
            id_to_output.get(a.get("output_id"), ""),
            a.get("name", ""),
            a.get("owner", ""),
            fmt_dd_mmm_yyyy(a.get("start")),
            fmt_dd_mmm_yyyy(a.get("end")),
            ",".join(a.get("kpi_ids") or []),  # machine IDs
            ", ".join(id_to_kpi.get(i, "") for i in (a.get("kpi_ids") or [])),  # display names
        ])

    # --- Budget (export, detailed) ---
    ws3 = wb.create_sheet("Budget")
    ws3.append([
        "Linked Output", "Linked Activity",
        "Budget Line Item", "Cost Category", "Sub Category",
        "Unit", "Unit Cost (USD)", "Quantity", "Total Cost (USD)"
    ])

    out_nums, _, act_nums = compute_numbers(include_activities=True)
    out_label = lambda oid: (f"Output {out_nums.get(oid, '')} — " + (
        next((o.get('name', '') for o in st.session_state.outputs if o['id'] == oid), ""))).strip(" —")
    act_lookup = activity_label_map(act_nums)

    def act_label(aid):
        return (act_lookup.get(aid, "") or "").strip(" —")

    for r in st.session_state.get("budget", []):
        ws3.append([
            out_label(r.get("output_id")) if r.get("output_id") else "",
            act_label(r.get("activity_id")) if r.get("activity_id") else "",
            r.get("item", ""),
            r.get("category", ""),
            r.get("subcategory", ""),
            r.get("unit", ""),
            float(r.get("unit_cost") or 0.0),
            float(r.get("qty") or 0.0),
            float(r.get("total_usd") or 0.0),
        ])

    # Number formats for currency/qty/total
    for row in ws3.iter_rows(min_row=2):
        row[6].number_format = '#,##0.00'  # unit cost
        row[7].number_format = '#,##0.00'  # quantity (allow decimals)
        row[8].number_format = '#,##0.00'  # total

    # --- Disbursement Schedule (export) ---
    wsd = wb.create_sheet("Disbursement Schedule")
    wsd.append([
        "KPIID",  # machine id for mapping back
        "Anticipated deliverable date",
        "Deliverable",
        "Maximum Grant instalment payable on satisfaction of this deliverable (USD)"
    ])

    from datetime import date as _date

    src = list(st.session_state.get("disbursement", [])) or [
        {
            "kpi_id": k["id"],
            "output_id": k.get("parent_id"),
            "anticipated_date": k.get("end_date") or k.get("start_date") or None,
            "deliverable": k.get("name", ""),
            "amount_usd": 0.0,
        }
        for k in st.session_state.kpis if bool(k.get("linked_payment"))
    ]

    rows = sorted(
        src,
        key=lambda d: (d.get("output_id"), d.get("anticipated_date") or _date(2100, 1, 1),
                       d.get("deliverable") or "")
    )

    for d in rows:
        # write actual date object; openpyxl will store a true Excel date
        wsd.append([
            d.get("kpi_id") or "",
            d.get("anticipated_date") or None,  # <-- date, not string
            (d.get("deliverable") or ""),
            float(d.get("amount_usd") or 0.0),
        ])

    # Number formats: Date (col B) and Amount (col D)
    for r in wsd.iter_rows(min_row=2):
        r[1].number_format = 'DD/mmm/YYYY'  # e.g., 01/Oct/2025
        r[3].number_format = '#,##0.00'

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    tabs[6].download_button(
        "📥 Download Backup File (Excel)",
        data=buf,
        file_name="Application_Submission.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

# --- Word export (Logframe as table)
if tabs[6].button("Generate Project Design Document (Word)"):
    try:
        word_buf = render_pdd()
        proj_title = (st.session_state.get("id_info", {}) or {}).get("title", "") or "Project"
        safe = re.sub(r"[^A-Za-z0-9]+", "_", proj_title).strip("_") or "Project"
        tabs[6].download_button(
            "📥 Download Project Design Document (Word)",
            data=word_buf,
            file_name=f"PDD_{safe}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except ModuleNotFoundError:
        tabs[6].error("`python-docx` is required. Install it with: pip install python-docx")
    except Exception as e:
        tabs[6].error(f"Could not generate the Word Document: {e}")
